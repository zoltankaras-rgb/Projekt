# office_handler.py
# Kancelária – kompletný handler (zladený a rozšírený)
# - Dashboardy, plánovanie, nákup surovín
# - Prehľady zásob (výrobný sklad, Sklad 2)
# - Katalóg produktov + recepty + krájané výrobky
# - B2C cenník, odmeny, objednávky
# - HACCP import/export + JSON úložisko (get_haccp_docs/get_haccp_doc_content/save_haccp_doc)
# - Dodávatelia (CRUD)
# - REPORT: Príjem z výroby (podľa dátumu PRÍJMU z Expedície) – bez uzávierky, smart fallback na cenu
# - Centrálne priemerné ceny: výrobné (vážený priemer), nákupné Sklad 2 (best-effort)
# - Stock helpers: create_production_item / update_production_item_qty / delete_production_item
#
# Bez ALTERov. Všade guardy a autodetekcia stĺpcov.

import os
import re
import json
import math
import uuid
import html
import unicodedata
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, date
from typing import Any, Dict, List, Optional, Tuple

from flask import make_response, request, send_file

import db_connector
import production_handler
import notification_handler
import b2b_handler  # používaš v app.py

COLL = 'utf8mb4_0900_ai_ci'   # konzistentná kolácia pre JOINy


# =================================================================
# === POMOCNÉ ======================================================
# =================================================================
def _fmt_time_hhmm(val):
    """Bezpečne prevedie datetime.time / datetime.timedelta na HH:MM; inak vráti ''. """
    if val is None:
        return ""
    # datetime.time
    try:
        return val.strftime("%H:%M")
    except Exception:
        pass
    # datetime.timedelta
    try:
        total = int(val.total_seconds())
        h = total // 3600
        m = (total % 3600) // 60
        return f"{h:02d}:{m:02d}"
    except Exception:
        return ""

def _columns(table: str) -> List[str]:
    try:
        rows = db_connector.execute_query("""
            SELECT COLUMN_NAME AS c
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_SCHEMA = DATABASE() AND TABLE_NAME=%s
        """, (table,)) or []
        return [r['c'] for r in rows]
    except Exception:
        return []

def _has_col(table: str, col: str) -> bool:
    try:
        r = db_connector.execute_query("""
            SELECT 1
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_SCHEMA = DATABASE() AND TABLE_NAME=%s AND COLUMN_NAME=%s
            LIMIT 1
        """, (table, col), fetch='one')
        return bool(r)
    except Exception:
        return False

def _zv_name_col() -> str:
    """V zaznamy_vyroba môže byť názov vo 'nazov_vyrobu' alebo 'nazov_vyrobku'."""
    try:
        r = db_connector.execute_query("""
            SELECT 1 FROM INFORMATION_SCHEMA.COLUMNS
             WHERE TABLE_SCHEMA = DATABASE() AND TABLE_NAME='zaznamy_vyroba' AND COLUMN_NAME='nazov_vyrobu'
             LIMIT 1
        """, fetch='one')
        return 'nazov_vyrobu' if r else 'nazov_vyrobku'
    except Exception:
        return 'nazov_vyrobku'

def _norm_key(s: str) -> str:
    if s is None:
        return ""
    t = unicodedata.normalize("NFKD", str(s))
    t = "".join(ch for ch in t if not unicodedata.combining(ch))
    return t.casefold().strip()

def _parse_num(x):
    if x is None:
        return None
    try:
        return float(str(x).replace(',', '.').strip())
    except Exception:
        return None

def _set_request_collation_slk():
    try:
        db_connector.execute_query("SET collation_connection = 'utf8mb4_slovak_ci'", fetch='none')
    except Exception:
        pass


# =================================================================
# === B2C – META (obrázky/popis) ==================================
# =================================================================

def _b2c_meta_path():
    base = os.path.dirname(__file__)
    folder = os.path.join(base, 'static', 'uploads', 'b2c')
    os.makedirs(folder, exist_ok=True)
    return os.path.join(folder, '_b2c_meta.json')

def _b2c_meta_load():
    try:
        with open(_b2c_meta_path(), 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}

def _b2c_meta_save(d: dict):
    p = _b2c_meta_path()
    tmp = p + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(d, f, ensure_ascii=False, indent=2)
    os.replace(tmp, p)

def _b2c_migrate_images_map():
    base = os.path.dirname(__file__)
    old = os.path.join(base, 'static', 'uploads', 'b2c', '_images_map.json')
    if not os.path.exists(old):
        return
    try:
        with open(old, 'r', encoding='utf-8') as f:
            old_map = json.load(f)
    except Exception:
        return
    meta = _b2c_meta_load()
    changed = False
    for ean, url in (old_map or {}).items():
        if not ean:
            continue
        meta.setdefault(ean, {})
        if url and meta[ean].get('obrazok') != url:
            meta[ean]['obrazok'] = url
            changed = True
    if changed:
        _b2c_meta_save(meta)
    try:
        os.remove(old)
    except Exception:
        pass

def upload_b2c_image():
    """Upload obrázka pre B2C cenník. Vráti URL do /static/uploads/b2c/."""
    file = request.files.get('file')
    if not file or not file.filename:
        return {"error": "Chýba súbor na nahratie."}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ('.jpg', '.jpeg', '.png', '.gif', '.webp'):
        return {"error": "Nepodporovaný formát obrázka."}
    fname = f"{uuid.uuid4().hex}{ext}"
    upload_dir = os.path.join(os.path.dirname(__file__), 'static', 'uploads', 'b2c')
    os.makedirs(upload_dir, exist_ok=True)
    file.save(os.path.join(upload_dir, fname))
    return {"url": f"/static/uploads/b2c/{fname}"}


# =================================================================
# === DASHBOARD KANCELÁRIA =========================================
# =================================================================

def get_kancelaria_dashboard_data():
    """Dashboard: suroviny pod minimom, finálny tovar pod minimom, akcie, top produkty, timeseries výroby."""
    # 1) suroviny pod minimom
    low_stock_raw = db_connector.execute_query("""
        SELECT nazov as name, mnozstvo as quantity, min_zasoba as minStock
        FROM sklad
        WHERE min_zasoba > 0 AND mnozstvo < min_zasoba
        ORDER BY nazov
    """) or []

    # 2) finálny tovar pod minimom
    all_goods = db_connector.execute_query("""
        SELECT
            nazov_vyrobku, predajna_kategoria, aktualny_sklad_finalny_kg,
            minimalna_zasoba_kg, minimalna_zasoba_ks, mj, vaha_balenia_g, typ_polozky
        FROM produkty
        WHERE typ_polozky = 'produkt' OR typ_polozky LIKE 'VÝROBOK%%' OR typ_polozky LIKE 'TOVAR%%'
    """) or []

    low_stock_goods_list = []
    for p in all_goods:
        stock_kg = float(p.get('aktualny_sklad_finalny_kg') or 0.0)
        min_stock_kg = float(p.get('minimalna_zasoba_kg') or 0.0)
        min_stock_ks = float(p.get('minimalna_zasoba_ks') or 0.0)
        mj = p.get('mj')
        weight_g = float(p.get('vaha_balenia_g') or 0.0)

        is_below_min = False
        if mj == 'ks':
            current_stock_ks = math.floor((stock_kg * 1000) / weight_g) if weight_g > 0 else 0
            if min_stock_ks > 0 and current_stock_ks < min_stock_ks:
                is_below_min = True
            current_stock_display = f"{current_stock_ks} ks"
            min_stock_display = f"{int(min_stock_ks)} ks"
        else:
            if min_stock_kg > 0 and stock_kg < min_stock_kg:
                is_below_min = True
            current_stock_display = f"{stock_kg:.2f} kg"
            min_stock_display = f"{min_stock_kg:.2f} kg"

        if is_below_min:
            low_stock_goods_list.append({
                "name": p['nazov_vyrobku'],
                "category": p.get('predajna_kategoria') or 'Nezaradené',
                "currentStock": current_stock_display,
                "minStock": min_stock_display
            })

    low_stock_goods_categorized = {}
    for item in low_stock_goods_list:
        low_stock_goods_categorized.setdefault(item['category'], []).append(item)

    # 3) aktívne akcie
    active_promos = db_connector.execute_query("""
        SELECT promo.product_name, promo.sale_price_net, promo.end_date, chain.name as chain_name
        FROM b2b_promotions promo
        JOIN b2b_retail_chains chain ON promo.chain_id = chain.id
        WHERE CURDATE() BETWEEN promo.start_date AND promo.end_date
        ORDER BY chain.name, promo.product_name
    """) or []

    # 4) TOP produkty + timeseries (opravené GROUP BY)
    top_products = db_connector.execute_query(f"""
        SELECT
            p.nazov_vyrobku AS name,
            SUM(COALESCE(zv.realne_mnozstvo_kg,0)) AS total
        FROM zaznamy_vyroba zv
        JOIN produkty p
          ON TRIM(zv.{_zv_name_col()}) = TRIM(p.nazov_vyrobku)
        WHERE zv.datum_ukoncenia >= CURDATE() - INTERVAL 30 DAY
          AND zv.stav IN ('Ukončené','Dokončené')
          AND COALESCE(zv.realne_mnozstvo_kg,0) > 0
        GROUP BY p.nazov_vyrobku
        ORDER BY total DESC
        LIMIT 5
    """) or []

    production_timeseries = db_connector.execute_query("""
        SELECT DATE_FORMAT(datum_ukoncenia, '%Y-%m-%d') as production_date,
               SUM(COALESCE(realne_mnozstvo_kg,0)) as total_kg
        FROM zaznamy_vyroba
        WHERE datum_ukoncenia >= CURDATE() - INTERVAL 30 DAY
          AND stav IN ('Ukončené','Dokončené')
        GROUP BY production_date
        ORDER BY production_date ASC
    """) or []

    return {
        "lowStockRaw": low_stock_raw,
        "lowStockGoods": low_stock_goods_categorized,
        "activePromotions": active_promos,
        "topProducts": top_products,
        "timeSeriesData": production_timeseries
    }

def get_kancelaria_base_data():
    products_list = db_connector.execute_query("""
        SELECT nazov_vyrobku
        FROM produkty
        WHERE (typ_polozky = 'produkt' OR TRIM(UPPER(typ_polozky)) LIKE 'VÝROBOK%%')
          AND nazov_vyrobku NOT IN (SELECT DISTINCT nazov_vyrobku FROM recepty)
        ORDER BY nazov_vyrobku
    """) or []
    categories_list = db_connector.execute_query("""
        SELECT DISTINCT kategoria_pre_recepty
        FROM produkty
        WHERE kategoria_pre_recepty IS NOT NULL AND kategoria_pre_recepty != ''
        ORDER BY kategoria_pre_recepty
    """) or []
    return {
        'warehouse': production_handler.get_warehouse_state(),
        'itemTypes': ['Mäso', 'Koreniny', 'Obaly - Črevá', 'Pomocný materiál'],
        'productsWithoutRecipe': [p['nazov_vyrobku'] for p in products_list],
        'recipeCategories': [c['kategoria_pre_recepty'] for c in categories_list]
    }


# =================================================================
# === EXPEDIČNÝ PLÁN / FORECAST ====================================
# =================================================================

def get_7_day_order_forecast():
    start_date = datetime.now().date()
    end_date   = start_date + timedelta(days=6)
    rows = db_connector.execute_query(f"""
        SELECT
            t.day AS pozadovany_datum_dodania,
            COALESCE(p.ean, t.ean) AS ean,
            COALESCE(p.nazov_vyrobku, CONCAT('EAN ', t.ean)) AS nazov_vyrobku,
            COALESCE(p.aktualny_sklad_finalny_kg, 0) AS aktualny_sklad_finalny_kg,
            COALESCE(p.mj, 'kg') AS mj,
            COALESCE(p.vaha_balenia_g, 0) AS vaha_balenia_g,
            COALESCE(p.predajna_kategoria, 'Nezaradené') AS predajna_kategoria,
            COALESCE(p.typ_polozky, '') AS typ_polozky,
            t.qty AS mnozstvo
        FROM (
            SELECT DATE(o.pozadovany_datum_dodania) AS day,
                   CONVERT(pol.ean_produktu USING utf8mb4) AS ean,
                   SUM(pol.mnozstvo) AS qty
            FROM b2b_objednavky_polozky pol
            JOIN b2b_objednavky o ON o.id = pol.objednavka_id
            WHERE DATE(o.pozadovany_datum_dodania) BETWEEN %s AND %s
            GROUP BY day, ean
        ) t
        LEFT JOIN produkty p
          ON  CONVERT(p.ean USING utf8mb4) COLLATE {COLL}
           =  t.ean COLLATE {COLL}
        ORDER BY t.day, nazov_vyrobku
    """, (start_date, end_date)) or []

    days = [(start_date + timedelta(days=i)).strftime("%Y-%m-%d") for i in range(7)]
    forecast: Dict[str, List[Dict[str, Any]]] = {}
    idx: Dict[tuple, Dict[str, Any]] = {}

    for r in rows:
        day = r["pozadovany_datum_dodania"].strftime("%Y-%m-%d")
        ean = (r["ean"] or "").strip()
        name = r["nazov_vyrobku"]
        mj   = r["mj"] or "kg"
        stock_kg = float(r.get("aktualny_sklad_finalny_kg") or 0.0)
        w_g     = float(r.get("vaha_balenia_g") or 0.0)
        stock_raw = math.floor((stock_kg * 1000) / w_g) if mj == "ks" and w_g > 0 else stock_kg
        stock_display = f"{int(stock_raw)} ks" if mj == "ks" else f"{stock_kg:.2f} kg"
        cat   = r.get("predajna_kategoria") or "Nezaradené"
        qty   = float(r.get("mnozstvo") or 0.0)
        manuf = (r.get("typ_polozky","").lower() == "produkt") or (r.get("typ_polozky","").upper().startswith("VÝROBOK"))

        k = (cat, ean)
        if k not in idx:
            item = {
                "name": name, "mj": mj,
                "stock_display": stock_display, "stock_raw": float(stock_raw),
                "total_needed": 0.0, "deficit": 0.0, "isManufacturable": bool(manuf),
                "daily_needs": {d: 0.0 for d in days},
            }
            idx[k] = item
            forecast.setdefault(cat, []).append(item)

        idx[k]["daily_needs"][day] = idx[k]["daily_needs"].get(day, 0.0) + qty
        idx[k]["total_needed"] += qty

    for items in forecast.values():
        for it in items:
            it["deficit"] = max(0.0, it["total_needed"] - float(it["stock_raw"]))

    for cat in list(forecast.keys()):
        forecast[cat].sort(key=lambda x: x["name"].casefold())

    return {"dates": days, "forecast": forecast}


def get_goods_purchase_suggestion():
    start_date = datetime.now().date()
    end_date   = start_date + timedelta(days=7)

    res = db_connector.execute_query(f"""
        SELECT COALESCE(p.ean, CONVERT(pol.ean_produktu USING utf8mb4)) AS ean,
               SUM(pol.mnozstvo) AS reserved_qty
        FROM b2b_objednavky_polozky pol
        JOIN b2b_objednavky o ON o.id = pol.objednavka_id
        LEFT JOIN produkty p
          ON  CONVERT(p.ean USING utf8mb4) COLLATE {COLL}
           =  CONVERT(pol.ean_produktu USING utf8mb4) COLLATE {COLL}
        WHERE DATE(o.pozadovany_datum_dodania) BETWEEN %s AND %s
        GROUP BY COALESCE(p.ean, CONVERT(pol.ean_produktu USING utf8mb4))
    """, (start_date, end_date)) or []
    rmap = { (r["ean"] or "").strip(): float(r.get("reserved_qty") or 0.0) for r in res }

    goods = db_connector.execute_query("""
        SELECT ean, nazov_vyrobku, COALESCE(aktualny_sklad_finalny_kg,0) AS stock_kg,
               COALESCE(minimalna_zasoba_kg,0) AS min_stock_kg, COALESCE(mj,'kg') AS mj
        FROM produkty
        WHERE TRIM(UPPER(typ_polozky)) IN ('PRODUKT','VÝROBOK','VÝROBOK_KUSOVY','VÝROBNÝ','VÝROBA','TOVAR','TOVAR_KUSOVY')
        ORDER BY nazov_vyrobku
    """) or []

    promo_set = {
        (r.get("product_ean") or "").strip()
        for r in (db_connector.execute_query("SELECT DISTINCT product_ean FROM b2b_promotions WHERE CURDATE() BETWEEN start_date AND end_date") or [])
    }

    out = []
    for g in goods:
        ean = (g["ean"] or "").strip()
        stock = float(g.get("stock_kg") or 0.0)
        min_kg = float(g.get("min_stock_kg") or 0.0)
        reserved = rmap.get(ean, 0.0)
        deficit = (min_kg + reserved) - stock
        if deficit > 0.0:
            out.append({
                "name": g["nazov_vyrobku"], "stock": round(stock, 2),
                "min_stock": round(min_kg, 2), "reserved": round(reserved, 2),
                "suggestion": round(deficit, 2), "unit": g.get("mj") or "kg",
                "is_promo": (ean in promo_set),
            })
    return out


# =================================================================
# === SKLAD – PREHĽADY A OPERÁCIE ==================================
# =================================================================

def get_raw_material_stock_overview():
    rows = db_connector.execute_query("""
        SELECT sv.nazov, COALESCE(sv.mnozstvo, 0) AS quantity,
               LOWER(COALESCE(s.typ, '')) AS typ, LOWER(COALESCE(s.podtyp, '')) AS podtyp
        FROM sklad_vyroba sv
        LEFT JOIN sklad s ON s.nazov = sv.nazov
        ORDER BY sv.nazov
    """) or []
    return {
        "items": [{"nazov": r["nazov"],
                   "quantity": float(r["quantity"] or 0),
                   "typ": r["typ"] or "", "podtyp": r["podtyp"] or ""} for r in rows]
    }

def get_production_stock_overview():
    rows = db_connector.execute_query("""
        SELECT sv.nazov, COALESCE(sv.mnozstvo, 0) AS mnozstvo,
               LOWER(COALESCE(s.typ, '')) AS typ, LOWER(COALESCE(s.podtyp, '')) AS podtyp
        FROM sklad_vyroba sv
        LEFT JOIN sklad s ON s.nazov = sv.nazov
        ORDER BY sv.nazov
    """) or []

    def resolve_cat(typ, podtyp):
        t = (typ or '').lower()
        p = (podtyp or '').lower()
        if t in ('maso','mäso') or p == 'maso': return 'maso'
        if t == 'koreniny' or p == 'koreniny': return 'koreniny'
        if t == 'obal': return 'obal'
        if t == 'pomocny_material': return 'pomocny_material'
        return 'nezaradene'

    items = []
    for r in rows:
        items.append({
            "nazov": r["nazov"],
            "mnozstvo": float(r["mnozstvo"] or 0),
            "typ": r["typ"] or "",
            "podtyp": r["podtyp"] or "",
            "cat": resolve_cat(r['typ'], r['podtyp'])
        })
    return {"items": items}

def receive_production_stock(payload):
    """
    Hromadný príjem do výrobného skladu (využíva aj /api/kancelaria/receiveStockItems).
    Použije vážený priemer do sklad.nakupna_cena (ak price je zadaná) a zaloguje do zaznamy_prijem.
    """
    if not isinstance(payload, dict):
        return {"error": "Neplatné dáta."}
    items = payload.get('items') or []
    if not items:
        return {"error": "Žiadne položky na príjem."}

    conn = db_connector.get_connection()
    try:
        cur = conn.cursor()
        for it in items:
            category = (it.get('category') or '').strip().lower()
            name     = (it.get('name') or '').strip()
            qty      = float(it.get('quantity') or 0)
            price    = it.get('price', None)
            note     = (it.get('note') or '').strip()
            dt_str   = it.get('date')
            when     = dt_str if dt_str else datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            if not name: return {"error": "Položka bez názvu."}
            if qty <= 0: return {"error": f"Nevyhovujúce množstvo pre '{name}'."}

            if category in ('maso','mäso'):
                src = (it.get('source') or '').strip().lower()
                if src not in ('rozrabka','expedicia','externy','ine'):
                    return {"error": f"Pre '{name}' (mäso) zvoľ Zdroj: rozrabka / expedicia / externy / ine."}
                prijem_typ = src
            else:
                prijem_typ = 'dodavatel'

            # karta v `sklad` musí existovať
            cur.execute("SELECT COALESCE(mnozstvo,0), COALESCE(nakupna_cena,0) FROM sklad WHERE nazov=%s FOR UPDATE", (name,))
            row = cur.fetchone()
            if row is None:
                return {"error": f"Položka '{name}' nie je založená v sklade. Najskôr ju vytvor v katalógu surovín."}
            central_qty, current_avg = float(row[0] or 0), float(row[1] or 0)

            # zásoba vo výrobnom sklade – pre váženie
            cur.execute("SELECT COALESCE(mnozstvo,0) FROM sklad_vyroba WHERE nazov=%s FOR UPDATE", (name,))
            r2 = cur.fetchone()
            prod_qty = float(r2[0]) if r2 is not None else 0.0

            # vážený priemer v sklad.nakupna_cena len ak price je zadaná
            if price is not None:
                total_before = central_qty + prod_qty
                new_total = total_before + qty
                new_avg = (current_avg * total_before + float(price) * qty) / new_total if new_total > 0 else float(price)
                cur.execute("UPDATE sklad SET nakupna_cena=%s WHERE nazov=%s", (new_avg, name))

            # navýš výrobný sklad
            cur.execute("""
                INSERT INTO sklad_vyroba (nazov, mnozstvo) VALUES (%s, %s)
                ON DUPLICATE KEY UPDATE mnozstvo = COALESCE(mnozstvo,0) + VALUES(mnozstvo)
            """, (name, qty))

            # log príjmu
            cur.execute("""
                INSERT INTO zaznamy_prijem (datum, nazov_suroviny, mnozstvo_kg, nakupna_cena_eur_kg, typ, poznamka_dodavatel)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (when, name, qty, price if price is not None else None, prijem_typ, note))

        conn.commit()
        return {"message": f"Prijatých {len(items)} položiek do výrobného skladu."}
    except Exception:
        if conn: conn.rollback()
        raise
    finally:
        if conn and conn.is_connected(): conn.close()

def receive_multiple_stock_items(payload: Dict[str, Any]):
    """Alias – back-compat s /api/kancelaria/receiveStockItems."""
    return receive_production_stock(payload)

def transfer_to_production(payload: Dict[str, Any], user: Optional[Dict[str, Any]] = None):
    """Rýchly transfer do výrobného skladu (navýšenie množstva)."""
    name = (payload or {}).get('name')
    qty  = _parse_num((payload or {}).get('quantity'))
    if not name or not qty or qty <= 0:
        return {"error": "Chýba názov alebo množstvo > 0."}
    db_connector.execute_query("""
        INSERT INTO sklad_vyroba (nazov, mnozstvo) VALUES (%s,%s)
        ON DUPLICATE KEY UPDATE mnozstvo = COALESCE(mnozstvo,0) + VALUES(mnozstvo)
    """, (name.strip(), float(qty)), fetch='none')
    return {"message": "Naskladnené do výroby."}

# ---- STOCK: production items (pre /api/kancelaria/stock/*) -------------------

def create_production_item(data: Dict[str, Any]):
    """
    Vytvorí (alebo aktualizuje) finálny produkt v `produkty`.
    Očakáva: { ean, nazov_vyrobku, mj('kg'|'ks'), vaha_balenia_g?, predajna_kategoria?, dph? }
    """
    ean  = (data or {}).get('ean')
    name = (data or {}).get('nazov_vyrobku') or (data or {}).get('name')
    mj   = (data or {}).get('mj') or 'kg'
    w_g  = data.get('vaha_balenia_g')
    cat  = (data or {}).get('predajna_kategoria') or 'Výrobky'
    dph  = data.get('dph', 19.0)

    if not ean or not name:
        return {"error": "Chýba EAN alebo názov produktu."}

    exists = db_connector.execute_query("SELECT ean FROM produkty WHERE ean=%s", (ean,), fetch='one')
    if exists:
        db_connector.execute_query("""
            UPDATE produkty SET nazov_vyrobku=%s, mj=%s, vaha_balenia_g=%s, predajna_kategoria=%s, dph=%s
            WHERE ean=%s
        """, (name, mj, (int(w_g) if w_g else None), cat, float(dph or 0), ean), fetch='none')
        return {"message": "Produkt aktualizovaný."}

    db_connector.execute_query("""
        INSERT INTO produkty (ean, nazov_vyrobku, mj, typ_polozky, vaha_balenia_g, predajna_kategoria, dph)
        VALUES (%s,%s,%s,'produkt',%s,%s,%s)
    """, (ean, name, mj, (int(w_g) if w_g else None), cat, float(dph or 0)), fetch='none')
    return {"message": "Produkt vytvorený."}

def update_production_item_qty(data: Dict[str, Any]):
    """
    Nastaví alebo upraví `aktualny_sklad_finalny_kg` pre produkt.
    Očakáva: { ean, quantity_kg, mode?="set"|"delta" } (default: "set")
    """
    ean = (data or {}).get('ean')
    qty = _parse_num((data or {}).get('quantity_kg'))
    mode = (data or {}).get('mode') or 'set'
    if not ean or qty is None:
        return {"error": "Chýba EAN alebo množstvo."}

    if mode == 'delta':
        db_connector.execute_query(
            "UPDATE produkty SET aktualny_sklad_finalny_kg = COALESCE(aktualny_sklad_finalny_kg,0) + %s WHERE ean=%s",
            (qty, ean), fetch='none'
        )
    else:
        db_connector.execute_query(
            "UPDATE produkty SET aktualny_sklad_finalny_kg = %s WHERE ean=%s",
            (qty, ean), fetch='none'
        )
    return {"message": "Množstvo upravené."}

def delete_production_item(data: Dict[str, Any]):
    """
    Bezpečné zmazanie produktu (ak nie je referencovaný).
    Očakáva: { ean }
    """
    ean = (data or {}).get('ean')
    if not ean:
        return {"error": "Chýba EAN."}
    used_r = db_connector.execute_query(
        "SELECT 1 FROM recepty WHERE nazov_vyrobku = (SELECT nazov_vyrobku FROM produkty WHERE ean=%s) LIMIT 1",
        (ean,), fetch='one'
    )
    used_s = db_connector.execute_query(
        "SELECT 1 FROM produkty WHERE zdrojovy_ean=%s LIMIT 1", (ean,), fetch='one'
    )
    if used_r or used_s:
        return {"error": "Nemožno vymazať – produkt je referencovaný (recept alebo krájanie)."}
    db_connector.execute_query("DELETE FROM produkty WHERE ean=%s", (ean,), fetch='none')
    return {"message": "Produkt zmazaný."}

# ---- add new stock item (raw warehouse) -------------------------------------

def add_new_stock_item(data: Dict[str, Any]):
    """
    Bez ALTER: bezpečne založí novú kartu v `sklad`.
    Očakáva minimálne: name (nazov).
    Voliteľne: ean, min_zasoba, default_cena_eur_kg, nakupna_cena, kategoria/typ/podtyp, is_infinite_stock.
    """
    name = (data or {}).get('name') or (data or {}).get('nazov')
    if not name:
        return {"error": "Chýba názov položky."}
    name = name.strip()
    cols = _columns('sklad')
    if not cols:
        return {"error": "Tabuľka 'sklad' neexistuje."}

    payload = {
        'nazov': name,
        'ean': (data or {}).get('ean'),
        'mnozstvo': float((data or {}).get('mnozstvo') or 0.0),
        'min_zasoba': (data or {}).get('min_zasoba'),
        'nakupna_cena': (data or {}).get('nakupna_cena'),
        'default_cena_eur_kg': (data or {}).get('default_cena_eur_kg'),
        'kategoria': (data or {}).get('kategoria') or (data or {}).get('typ') or (data or {}).get('podtyp'),
        'typ': (data or {}).get('typ'),
        'podtyp': (data or {}).get('podtyp'),
        'is_infinite_stock': 1 if (data or {}).get('is_infinite_stock') else 0
    }
    use_cols = [c for c in payload.keys() if c in cols and payload[c] is not None]
    values   = [payload[c] for c in use_cols]
    placeholders = ",".join(["%s"]*len(use_cols))

    exists = db_connector.execute_query("SELECT nazov FROM sklad WHERE nazov=%s LIMIT 1", (name,), fetch='one')
    if exists:
        set_list = ", ".join([f"{c}=%s" for c in use_cols if c != 'nazov'])
        params   = [payload[c] for c in use_cols if c != 'nazov'] + [name]
        db_connector.execute_query(f"UPDATE sklad SET {set_list} WHERE nazov=%s", tuple(params), fetch='none')
        return {"message": f"Položka '{name}' aktualizovaná."}

    cols_sql = ", ".join(use_cols)
    db_connector.execute_query(f"INSERT INTO sklad ({cols_sql}) VALUES ({placeholders})", tuple(values), fetch='none')
    return {"message": f"Položka '{name}' založená."}


# =================================================================
# === PRIEMERY CIEN – výrobná & nákupná (centrálne API) ============
# =================================================================

def _avg_purchase_costs_map_by_ean() -> dict:
    """Best-effort priemerná nákupná cena v sklade 2 (EAN → €/kg alebo €/ks podľa zdroja)."""
    candidates = [
        ('produkty_prijmy',  'ean',          'mnozstvo_kg', 'cena_eur_kg'),
        ('prijmy_sklad2',    'ean',          'mnozstvo_kg', 'cena_kg'),
        ('tovar_prijmy',     'ean_produktu', 'mnozstvo',    'cena_za_jednotku'),
        ('prijmy_produktov', 'ean',          'mnozstvo',    'cena'),
    ]
    out = {}
    for tbl, ce, cq, cp in candidates:
        if _has_col(tbl, ce) and _has_col(tbl, cq) and _has_col(tbl, cp):
            try:
                rows = db_connector.execute_query(
                    f"SELECT {ce} AS ean, SUM({cq}) AS qty, SUM({cq}*{cp}) AS val FROM {tbl} GROUP BY {ce}"
                ) or []
                for r in rows:
                    q = float(r['qty'] or 0.0); v = float(r['val'] or 0.0)
                    if q > 0: out[(r['ean'] or '').strip()] = v / q
                if out: break
            except Exception:
                continue
    return out

def get_avg_costs_catalog():
    """
    Centrálne priemery pre Kanceláriu:
     - avg_manufacturing_unit_cost: vážený priemer z 'zaznamy_vyroba.cena_za_jednotku' (€/kg alebo €/ks podľa MJ)
     - avg_purchase_unit_cost: best-effort z príjmov Sklad 2 (ak nie je zdroj, vráti None)
    """
    zv = _zv_name_col()
    rows = db_connector.execute_query(f"""
        SELECT p.ean, TRIM(p.nazov_vyrobku) AS product_name, p.mj AS prod_mj,
               SUM((CASE WHEN p.mj='kg' THEN COALESCE(zv.realne_mnozstvo_kg,0)
                         ELSE COALESCE(zv.realne_mnozstvo_ks,0) END)
                   * COALESCE(zv.cena_za_jednotku,0)) AS sum_cost_units,
               SUM(CASE WHEN p.mj='kg' THEN COALESCE(zv.realne_mnozstvo_kg,0)
                        ELSE COALESCE(zv.realne_mnozstvo_ks,0) END) AS sum_units
        FROM zaznamy_vyroba zv
        JOIN produkty p ON TRIM(zv.{zv}) = TRIM(p.nazov_vyrobku)
        WHERE COALESCE(zv.cena_za_jednotku,0) > 0
        GROUP BY p.ean, TRIM(p.nazov_vyrobku), p.mj
    """) or []
    out = []
    for r in rows:
        units = float(r['sum_units'] or 0.0)
        avg   = (float(r['sum_cost_units'] or 0.0) / units) if units > 0 else 0.0
        out.append({
            "ean": (r['ean'] or '').strip(),
            "product": r['product_name'],
            "unit": r['prod_mj'],
            "avg_manufacturing_unit_cost": round(avg, 4)
        })

    purchase_by_ean = _avg_purchase_costs_map_by_ean()
    for rec in out:
        rec["avg_purchase_unit_cost"] = (None if rec["ean"] not in purchase_by_ean
                                         else round(float(purchase_by_ean[rec["ean"]]), 4))
    return {"rows": out}


# =================================================================
# === PREHĽAD CENTRÁLNEHO KATALÓGU (Sklad 2) =======================
# =================================================================

def get_comprehensive_stock_view():
    """Prehľad finálnych produktov (centrálny sklad) + priemery cien (výrobné/nákupné)."""
    q = """
        SELECT
            p.ean, p.nazov_vyrobku AS name, p.predajna_kategoria AS category,
            p.aktualny_sklad_finalny_kg AS stock_kg, p.vaha_balenia_g, p.mj AS unit,
            (
              SELECT ROUND(zv.celkova_cena_surovin / NULLIF(zv.realne_mnozstvo_kg, 0), 4)
              FROM zaznamy_vyroba zv
              WHERE (zv.nazov_vyrobu = p.nazov_vyrobku OR zv.nazov_vyrobku = p.nazov_vyrobku)
                AND zv.celkova_cena_surovin IS NOT NULL
                AND zv.realne_mnozstvo_kg IS NOT NULL
              ORDER BY COALESCE(zv.datum_ukoncenia, zv.datum_vyroby) DESC
              LIMIT 1
            ) AS last_cost_per_kg
        FROM produkty p
        WHERE p.typ_polozky = 'produkt' OR p.typ_polozky LIKE 'VÝROBOK%%' OR p.typ_polozky LIKE 'TOVAR%%'
        ORDER BY category, name
    """
    rows = db_connector.execute_query(q) or []

    # centrálne výrobné priemery (€/kg alebo €/ks podľa MJ) z dávok s cenou
    zv = _zv_name_col()
    mc_rows = db_connector.execute_query(f"""
        SELECT TRIM(p.nazov_vyrobku) AS pn, p.mj AS mj,
               SUM((CASE WHEN p.mj='kg' THEN COALESCE(zv.realne_mnozstvo_kg,0)
                         ELSE COALESCE(zv.realne_mnozstvo_ks,0) END)
                   * COALESCE(zv.cena_za_jednotku,0)) AS sum_cost_units,
               SUM(CASE WHEN p.mj='kg' THEN COALESCE(zv.realne_mnozstvo_kg,0)
                        ELSE COALESCE(zv.realne_mnozstvo_ks,0) END) AS sum_units
          FROM zaznamy_vyroba zv
          JOIN produkty p ON TRIM(zv.{zv}) = TRIM(p.nazov_vyrobku)
         WHERE COALESCE(zv.cena_za_jednotku,0) > 0
         GROUP BY TRIM(p.nazov_vyrobku), p.mj
    """) or []
    manuf_index = {}
    for r in mc_rows:
        su = float(r['sum_units'] or 0.0)
        avg = (float(r['sum_cost_units'] or 0.0)/su) if su>0 else 0.0
        manuf_index[(r['pn'], r['mj'])] = avg

    # nákupné best-effort cez EAN
    purchase_by_ean = _avg_purchase_costs_map_by_ean()

    grouped: Dict[str, List[Dict[str, Any]]] = {}
    flat: List[Dict[str, Any]] = []

    for p in rows:
        unit = p.get('unit') or 'kg'
        qty_kg = float(p.get('stock_kg') or 0.0)
        w = float(p.get('vaha_balenia_g') or 0.0)
        qty = (qty_kg * 1000 / w) if unit == 'ks' and w > 0 else qty_kg

        manuf_avg = manuf_index.get((p['name'], unit), 0.0)
        purchase_avg = None
        if p['ean']:
            base = purchase_by_ean.get((p['ean'] or '').strip())
            if base is not None:
                purchase_avg = base
                if unit == 'ks' and w > 0:
                    purchase_avg = base * (w/1000.0)

        item = {
            "ean": p['ean'],
            "name": p['name'],
            "category": p.get('category') or 'Nezaradené',
            "quantity": qty,
            "unit": unit,
            "sklad1": 0.0,
            "sklad2": qty_kg,
            "last_cost_per_kg": float(p.get('last_cost_per_kg') or 0.0),
            "avg_manufacturing_unit_cost": round(float(manuf_avg), 4),
            "avg_purchase_unit_cost": (None if purchase_avg is None else round(float(purchase_avg), 4)),
        }
        flat.append(item)
        grouped.setdefault(item['category'], []).append(item)

    return {"products": flat, "groupedByCategory": grouped}


# =================================================================
# === CENNÍKY / PROMO (predvolené ceny) ============================
# =================================================================

def _avg_costs_index_by_ean() -> Dict[str, Dict[str, Any]]:
    idx = {}
    rows = get_avg_costs_catalog().get("rows") or []
    for r in rows:
        if r["ean"]:
            idx[r["ean"]] = r
    return idx

def get_promotions_data():
    chains = db_connector.execute_query("SELECT * FROM b2b_retail_chains ORDER BY name") or []
    promos = db_connector.execute_query("""
        SELECT p.*, c.name as chain_name
        FROM b2b_promotions p
        JOIN b2b_retail_chains c ON p.chain_id = c.id
        ORDER BY p.start_date DESC
    """) or []
    products = db_connector.execute_query("""
        SELECT ean, nazov_vyrobku as name, mj
        FROM produkty
        WHERE typ_polozky = 'produkt' OR typ_polozky LIKE 'TOVAR%%'
        ORDER BY name
    """) or []

    avg_idx = _avg_costs_index_by_ean()
    for p in products:
        rec = avg_idx.get((p.get("ean") or "").strip())
        p["default_unit_cost"] = (rec or {}).get("avg_manufacturing_unit_cost")

    return {"chains": chains, "promotions": promos, "products": products}

def manage_promotion_chain(data: Dict[str, Any]):
    action = (data or {}).get('action')
    if action == 'add':
        name = (data.get('name') or '').strip()
        if not name: return {"error": "Názov reťazca je povinný."}
        try:
            db_connector.execute_query("INSERT INTO b2b_retail_chains (name) VALUES (%s)", (name,), fetch='none')
            return {"message": "Reťazec pridaný."}
        except Exception:
            return {"error": "Reťazec s týmto názvom už existuje."}
    elif action == 'delete':
        chain_id = data.get('id')
        if not chain_id: return {"error": "Chýba ID reťazca."}
        db_connector.execute_query("DELETE FROM b2b_retail_chains WHERE id = %s", (chain_id,), fetch='none')
        return {"message": "Reťazec vymazaný."}
    return {"error": "Neznáma akcia."}

def _get_product_name_by_ean(ean: str) -> str:
    row = db_connector.execute_query("SELECT nazov_vyrobku AS name FROM produkty WHERE ean=%s LIMIT 1", (ean,), fetch='one')
    if row and row.get('name'): return row['name']
    row = db_connector.execute_query("SELECT nazov AS name FROM sklad WHERE ean=%s LIMIT 1", (ean,), fetch='one')
    if row and row.get('name'): return row['name']
    return ean

def _get_chain_name(chain_id: Optional[int]) -> Optional[str]:
    if not chain_id:
        return None
    row = db_connector.execute_query("SELECT name FROM b2b_retail_chains WHERE id=%s LIMIT 1", (chain_id,), fetch='one')
    if row and row.get('name'): return row['name']
    row = db_connector.execute_query("SELECT nazov AS name FROM b2b_retail_chains WHERE id=%s LIMIT 1", (chain_id,), fetch='one')
    if row and row.get('name'): return row['name']
    return None

def save_promotion(data: Dict[str, Any]):
    if not data:
        return {"error": "Chýbajú údaje."}
    for key in ("chain_id", "ean", "start_date", "end_date", "sale_price_net"):
        if data.get(key) in (None, "", []):
            return {"error": f"Chýba pole: {key}"}

    ean = (str(data["ean"]).strip())
    product_name = _get_product_name_by_ean(ean)
    try:
        price_net = float(data["sale_price_net"])
    except Exception:
        return {"error": "sale_price_net musí byť číslo."}

    chain_id_int = None
    try:
        if str(data.get("chain_id")).strip() != "":
            chain_id_int = int(data.get("chain_id"))
    except Exception:
        chain_id_int = None

    chain_name = _get_chain_name(chain_id_int)
    start_date = (str(data["start_date"]) or "")[:10]
    end_date   = (str(data["end_date"]) or "")[:10]
    promo_name = f"{product_name} {start_date}–{end_date}"
    if chain_name:
        promo_name = f"{chain_name} – {promo_name}"

    cols = ["chain_id", "name", "product_ean", "product_name", "start_date", "end_date", "sale_price_net"]
    vals = [chain_id_int, promo_name, ean, product_name, start_date, end_date, price_net]
    placeholders = ["%s"] * len(cols)

    try:
        has_created_at = db_connector.execute_query("""
            SELECT 1 FROM INFORMATION_SCHEMA.COLUMNS
             WHERE TABLE_SCHEMA = DATABASE()
               AND TABLE_NAME = 'b2b_promotions'
               AND COLUMN_NAME = 'created_at'
             LIMIT 1
        """, fetch='one')
    except Exception:
        has_created_at = None

    if has_created_at:
        cols.append("created_at"); vals.append(datetime.now()); placeholders.append("%s")

    sql = f"INSERT INTO b2b_promotions ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
    db_connector.execute_query(sql, tuple(vals), fetch='none')
    return {"ok": True, "message": "Akcia bola uložená."}

def delete_promotion(data: Dict[str, Any]):
    promo_id = data.get('id')
    if not promo_id: return {"error": "Chýba ID akcie."}
    db_connector.execute_query("DELETE FROM b2b_promotions WHERE id = %s", (promo_id,), fetch='none')
    return {"message": "Akcia bola vymazaná."}


# =================================================================
# === KATALÓG – CRUD PRODUKTOV =====================================
# =================================================================

def get_catalog_management_data():
    products = db_connector.execute_query("""
        SELECT ean, nazov_vyrobku, typ_polozky, kategoria_pre_recepty, predajna_kategoria, dph
        FROM produkty
        ORDER BY typ_polozky, nazov_vyrobku
    """) or []
    recipe_categories = [r['kategoria_pre_recepty'] for r in (db_connector.execute_query("""
        SELECT DISTINCT kategoria_pre_recepty
        FROM produkty
        WHERE kategoria_pre_recepty IS NOT NULL AND kategoria_pre_recepty != ''
        ORDER BY 1
    """) or [])]
    sale_categories = ['Výrobky','Bravčové mäso chladené','Bravčové mäso mrazené','Hovädzie mäso chladené',
                       'Hovädzie mäso mrazené','Hydinové mäso chladené','Hydinové mäso mrazené',
                       'Ryby mrazené','Zelenina','Tovar']
    item_types = ['VÝROBOK','VÝROBOK_KRAJANY','VÝROBOK_KUSOVY','TOVAR','TOVAR_KUSOVY']
    dph_rates = [5.00, 10.00, 19.00, 23.00]
    return {"products": products, "recipe_categories": recipe_categories,
            "sale_categories": sale_categories, "item_types": item_types, "dph_rates": dph_rates}

def add_catalog_item(data):
    ean  = (data.get('new_catalog_ean') or '').strip()
    name = (data.get('new_catalog_name') or '').strip()
    item_type = (data.get('new_catalog_item_type') or '').strip()
    dph  = data.get('new_catalog_dph')
    sale_cat = (data.get('new_catalog_sale_category') or None)
    if not all([ean, name, item_type, dph]): return {"error": "EAN, Názov, Typ a DPH sú povinné."}
    if db_connector.execute_query("SELECT ean FROM produkty WHERE ean=%s", (ean,), fetch='one'):
        return {"error": f"EAN '{ean}' už existuje."}
    if db_connector.execute_query("SELECT ean FROM produkty WHERE TRIM(UPPER(nazov_vyrobku))=TRIM(UPPER(%s))", (name,), fetch='one'):
        return {"error": f"Názov '{name}' už existuje."}
    mj = 'ks' if item_type in ['VÝROBOK_KRAJANY','VÝROBOK_KUSOVY','TOVAR_KUSOVY'] else 'kg'
    db_connector.execute_query("""
        INSERT INTO produkty (ean, nazov_vyrobku, typ_polozky, kategoria_pre_recepty, predajna_kategoria, dph, mj)
        VALUES (%s,%s,%s,%s,%s,%s,%s)
    """, (ean, name, item_type, None, sale_cat, dph, mj), fetch='none')
    return {"message": f"Položka '{name}' pridaná."}

def update_catalog_item(data):
    ean = (data.get('ean') or '').strip()
    if not ean: return {"error": "Chýba EAN."}
    typ = (data.get('typ_polozky') or '').strip() or None
    db_connector.execute_query("""
        UPDATE produkty
           SET nazov_vyrobku=%s, typ_polozky=%s, kategoria_pre_recepty=%s, predajna_kategoria=%s, dph=%s
         WHERE ean=%s
    """, (data.get('nazov_vyrobku'), typ,
          data.get('kategoria_pre_recepty') or None,
          data.get('predajna_kategoria') or None,
          float(data.get('dph', 0) or 0), ean), fetch='none')
    return {"message": f"Položka {ean} aktualizovaná."}

def delete_catalog_item(data):
    ean = (data.get('ean') or '').strip()
    if not ean: return {"error": "Chýba EAN."}
    if db_connector.execute_query(
        "SELECT 1 FROM recepty WHERE nazov_vyrobku = (SELECT nazov_vyrobku FROM produkty WHERE ean=%s) LIMIT 1",
        (ean,), fetch='one'):
        return {"error": "Nemožno vymazať, je použitá v recepte."}
    if db_connector.execute_query("SELECT 1 FROM produkty WHERE zdrojovy_ean=%s LIMIT 1", (ean,), fetch='one'):
        return {"error": "Nemožno vymazať, je zdrojom pre krájanie."}
    db_connector.execute_query("DELETE FROM produkty WHERE ean=%s", (ean,), fetch='none')
    return {"message": f"Položka {ean} vymazaná."}


# =================================================================
# === RECEPTY ======================================================
# =================================================================

def add_new_recipe(recipe_data):
    product_name = (recipe_data or {}).get('productName')
    ingredients = (recipe_data or {}).get('ingredients') or []
    category = ((recipe_data or {}).get('newCategory') or '').strip() or (recipe_data or {}).get('category')
    if not product_name or not category:
        return {"error": "Chýba produkt alebo kategória."}
    rows_to_insert = []
    for ing in ingredients:
        name = (ing or {}).get('name'); qty = _parse_num((ing or {}).get('quantity'))
        if name and qty and qty > 0: rows_to_insert.append((product_name, name, qty))
    exists = db_connector.execute_query(
        "SELECT 1 FROM recepty WHERE TRIM(nazov_vyrobku)=TRIM(%s) LIMIT 1", (product_name,), fetch='one')
    if exists: return {"error": f"Recept pre '{product_name}' už existuje."}

    conn = db_connector.get_connection()
    try:
        cur = conn.cursor()
        if rows_to_insert:
            cur.executemany(
                "INSERT INTO recepty (nazov_vyrobku, nazov_suroviny, mnozstvo_na_davku_kg) VALUES (%s,%s,%s)",
                rows_to_insert
            )
        cur.execute("UPDATE produkty SET kategoria_pre_recepty=%s WHERE TRIM(nazov_vyrobku)=TRIM(%s)",
                    (category, product_name))
        conn.commit()
        return {"message": f"Recept pre '{product_name}' bol vytvorený."}
    except Exception:
        if conn: conn.rollback(); raise
    finally:
        if conn and conn.is_connected(): conn.close()

def get_all_recipes_for_editing():
    rows = db_connector.execute_query("""
      SELECT p.nazov_vyrobku, p.kategoria_pre_recepty
        FROM produkty p
        JOIN (SELECT DISTINCT TRIM(nazov_vyrobku) AS nazov_vyrobku FROM recepty) r
          ON TRIM(p.nazov_vyrobku)=r.nazov_vyrobku
       WHERE p.typ_polozky LIKE 'VÝROBOK%%'
       ORDER BY p.kategoria_pre_recepty, p.nazov_vyrobku
    """) or []
    out: Dict[str, List[str]] = {}
    for r in rows:
        cat = r.get('kategoria_pre_recepty') or 'Nezaradené'
        out.setdefault(cat, []).append(r['nazov_vyrobku'])
    return out

def get_recipe_details(product_name: str):
    if not product_name: return {"error": "Chýba názov produktu."}
    prod = db_connector.execute_query(
        "SELECT kategoria_pre_recepty FROM produkty WHERE TRIM(nazov_vyrobku)=TRIM(%s)",
        (product_name,), fetch='one') or {}
    category = (prod or {}).get('kategoria_pre_recepty') or ''

    rows = db_connector.execute_query(f"""
        SELECT r.nazov_suroviny AS name, r.mnozstvo_na_davku_kg AS quantity,
            (SELECT z.nakupna_cena_eur_kg
               FROM zaznamy_prijem z
              WHERE z.nazov_suroviny COLLATE {COLL} = r.nazov_suroviny COLLATE {COLL}
                AND z.nakupna_cena_eur_kg IS NOT NULL
              ORDER BY z.datum DESC LIMIT 1) AS last_price,
            s.is_infinite_stock, s.default_cena_eur_kg
        FROM recepty r
        LEFT JOIN sklad s ON s.nazov COLLATE {COLL} = r.nazov_suroviny COLLATE {COLL}
        WHERE TRIM(r.nazov_vyrobku)=TRIM(%s)
        ORDER BY r.nazov_suroviny
    """, (product_name,), fetch='all') or []

    ingredients = []
    for rec in rows:
        price = rec.get("last_price")
        if price is None: price = rec.get("default_cena_eur_kg")
        if price is None and int(rec.get("is_infinite_stock") or 0) == 1:
            price = 0.20  # voda/ľad/ovar
        ingredients.append({
            "name": rec["name"], "quantity": float(rec["quantity"]) if rec["quantity"] is not None else None,
            "last_price": float(price) if price is not None else None
        })
    return {"productName": product_name, "category": category, "ingredients": ingredients}

def update_recipe(recipe_data):
    product_name = (recipe_data or {}).get('productName')
    ingredients = (recipe_data or {}).get('ingredients') or []
    category = ((recipe_data or {}).get('newCategory') or '').strip() or (recipe_data or {}).get('category')
    if not product_name: return {"error": "Chýba názov produktu."}
    rows_to_insert = []
    for ing in ingredients:
        name = (ing or {}).get('name'); qty = _parse_num((ing or {}).get('quantity'))
        if name and qty and qty > 0: rows_to_insert.append((product_name, name, qty))

    conn = db_connector.get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM recepty WHERE TRIM(nazov_vyrobku)=TRIM(%s)", (product_name,))
        if rows_to_insert:
            cur.executemany(
                "INSERT INTO recepty (nazov_vyrobku, nazov_suroviny, mnozstvo_na_davku_kg) VALUES (%s,%s,%s)",
                rows_to_insert
            )
        if category:
            cur.execute("UPDATE produkty SET kategoria_pre_recepty=%s WHERE TRIM(nazov_vyrobku)=TRIM(%s)",
                        (category, product_name))
        conn.commit()
        return {"message": f"Recept pre '{product_name}' bol upravený."}
    except Exception:
        if conn: conn.rollback(); raise
    finally:
        if conn and conn.is_connected(): conn.close()

def delete_recipe(product_name):
    if not product_name: return {"error": "Chýba názov produktu."}
    db_connector.execute_query("DELETE FROM recepty WHERE TRIM(nazov_vyrobku)=TRIM(%s)", (product_name,), fetch='none')
    return {"message": f"Recept pre '{product_name}' vymazaný."}


# =================================================================
# === KRÁJANÉ PRODUKTY =============================================
# =================================================================

def get_slicing_management_data():
    sources = db_connector.execute_query("""
        SELECT ean, nazov_vyrobku AS name
        FROM produkty
        WHERE TRIM(UPPER(typ_polozky)) IN ('VÝROBOK','VÝROBOK_KUSOVY','PRODUKT')
        ORDER BY nazov_vyrobku
    """) or []
    sliced = db_connector.execute_query("""
        SELECT ean, nazov_vyrobku AS name, zdrojovy_ean, vaha_balenia_g
        FROM produkty
        WHERE TRIM(UPPER(typ_polozky)) = 'VÝROBOK_KRAJANY'
        ORDER BY nazov_vyrobku
    """) or []
    return {"sourceProducts": sources, "slicedProducts": sliced}

def link_sliced_product(data):
    source_ean = (data or {}).get('sourceEan')
    target_ean = (data or {}).get('targetEan')
    weight_raw = (data or {}).get('weight') or (data or {}).get('weight_g') or (data or {}).get('vaha')
    if not source_ean or not target_ean: return {"error": "Chýba EAN zdroja alebo cieľa."}
    try:
        w = int(round(float(str(weight_raw).strip())))
    except Exception:
        w = 0
    if w <= 0: return {"error": "Zadajte váhu balíčka v gramoch (> 0)."}

    db_connector.execute_query("""
        UPDATE produkty
           SET zdrojovy_ean=%s, typ_polozky='VÝROBOK_KRAJANY', vaha_balenia_g=%s
         WHERE ean=%s
    """, (source_ean, w, target_ean), fetch='none')
    return {"message": "Prepojené.", "savedWeight": w}

def create_and_link_sliced_product(data):
    source_ean = (data.get('sourceEan') or '').strip()
    new_name   = (data.get('name') or '').strip()
    new_ean    = (data.get('ean') or '').strip()
    new_weight = data.get('weight')
    if not all([source_ean, new_name, new_ean, new_weight]):
        return {"error": "Všetky polia sú povinné."}
    if db_connector.execute_query("SELECT ean FROM produkty WHERE ean=%s", (new_ean,), fetch='one'):
        return {"error": f"EAN '{new_ean}' už existuje."}
    src = db_connector.execute_query("SELECT predajna_kategoria, dph FROM produkty WHERE ean=%s", (source_ean,), fetch='one')
    if not src: return {"error": "Zdrojový produkt nebol nájdený."}
    sale_cat, dph_rate = src.get('predajna_kategoria', 'Výrobky'), src.get('dph', 19.00)

    db_connector.execute_query("""
        INSERT INTO produkty (ean, nazov_vyrobku, mj, typ_polozky, vaha_balenia_g, zdrojovy_ean, dph, predajna_kategoria)
        VALUES (%s,%s,'ks','VÝROBOK_KRAJANY',%s,%s,%s,%s)
    """, (new_ean, new_name, float(new_weight), source_ean, dph_rate, sale_cat), fetch='none')
    return {"message": f"Produkt '{new_name}' vytvorený a prepojený."}


# =================================================================
# === MINIMÁLNE ZÁSOBY (KATALÓG) ===================================
# =================================================================

def get_products_for_min_stock():
    return db_connector.execute_query("""
        SELECT ean, nazov_vyrobku as name, mj, minimalna_zasoba_kg as minStockKg, minimalna_zasoba_ks as minStockKs
        FROM produkty
        WHERE typ_polozky='produkt' OR typ_polozky IN ('VÝROBOK','VÝROBOK_KRAJANY','VÝROBOK_KUSOVY')
        ORDER BY nazov_vyrobku
    """) or []

def update_min_stock_levels(products_data):
    if not products_data:
        return {"error": "Žiadne dáta na aktualizáciu."}
    updates = []
    for p in products_data:
        ean = (p.get('ean') or '').strip()
        if not ean: continue
        kg = p.get('minStockKg'); ks = p.get('minStockKs')
        if kg in (None, ''): kg_val = None
        else:
            try: kg_val = float(str(kg).replace(',', '.'))
            except Exception: kg_val = None
        if ks in (None, ''): ks_val = None
        else:
            try: ks_val = int(float(str(ks).replace(',', '.')))
            except Exception: ks_val = None
        updates.append((kg_val, ks_val, ean))
    if not updates:
        return {"error": "Žiadne platné dáta na aktualizáciu."}
    db_connector.execute_query("""
        UPDATE produkty SET minimalna_zasoba_kg=%s, minimalna_zasoba_ks=%s WHERE ean=%s
    """, updates, fetch='none', multi=True)
    return {"message": f"Minimálne zásoby aktualizované pre {len(updates)} produktov."}


# =================================================================
# === REPORTY – štatistiky, príjem, inventúra, príjem podľa dátumu ==
# =================================================================

def get_production_stats(period, category):
    """
    Štatistiky výroby a škôd – bez textových porovnaní v SQL,
    s cenovým fallbackom na celkové náklady / reálne kg, ak chýba cena_za_jednotku.
    """
    now = datetime.now()
    p = (str(period) if period is not None else "").lower()
    if p.startswith("week"):
        start_dt = (now - timedelta(days=now.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
    elif p.startswith("month"):
        start_dt = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    else:
        start_dt = datetime(1970, 1, 1)

    cat_raw  = (str(category) if category is not None else "").strip()
    cat_norm = cat_raw.casefold()
    use_cat  = bool(cat_norm) and cat_norm not in {"všetky","vsetky","all"}

    prod_map_rows = db_connector.execute_query(
        "SELECT nazov_vyrobku, kategoria_pre_recepty, mj FROM produkty"
    ) or []
    name2info = {
        _norm_key(r.get("nazov_vyrobku")): {
            "kategoria": (r.get("kategoria_pre_recepty") or "Nezaradené"),
            "unit":      (r.get("mj") or "kg"),
        }
        for r in prod_map_rows
    }

    prod_rows = db_connector.execute_query("""
        SELECT *
          FROM zaznamy_vyroba
         WHERE stav IN ('Ukončené','Dokončené')
           AND datum_ukoncenia >= %s
         ORDER BY datum_ukoncenia DESC, id_davky DESC
    """, (start_dt,)) or []

    enriched_prod = []
    for r in prod_rows:
        key  = _norm_key(r.get("nazov_vyrobku"))
        meta = name2info.get(key, {"kategoria":"Nezaradené","unit": (r.get("mj") or "kg")})
        if use_cat and meta["kategoria"].casefold() != cat_norm:
            continue

        plan_kg = float(r.get("planovane_mnozstvo_kg") or 0.0)
        real_kg = float(r.get("realne_mnozstvo_kg") or 0.0)
        yield_pct = ((real_kg / plan_kg) * 100.0 - 100.0) if plan_kg > 0 else 0.0

        unit_cost = float(r.get("cena_za_jednotku") or 0.0)
        if unit_cost == 0.0 and real_kg > 0:
            total_cost = float(r.get("celkova_cena_surovin") or 0.0)
            if total_cost:
                unit_cost = total_cost / real_kg

        row = dict(r)
        row["kategoria_pre_recepty"] = meta["kategoria"]
        row["unit"]                  = meta["unit"]
        row["vytaznost"]             = yield_pct
        row["cena_bez_energii"]      = unit_cost
        row["cena_s_energiami"]      = (unit_cost * 1.15) if unit_cost else 0.0
        enriched_prod.append(row)

    dmg_rows = db_connector.execute_query("""
        SELECT *
          FROM skody
         WHERE datum >= %s
         ORDER BY datum DESC, id DESC
    """, (start_dt,)) or []

    dmg_ids = sorted({ d["id_davky"] for d in dmg_rows if d.get("id_davky") })
    zv_map  = {}
    if dmg_ids:
        placeholders = ",".join(["%s"] * len(dmg_ids))
        zvr = db_connector.execute_query(
            f"SELECT id_davky, nazov_vyrobku, celkova_cena_surovin FROM zaznamy_vyroba WHERE id_davky IN ({placeholders})",
            tuple(dmg_ids)
        ) or []
        for z in zvr:
            zv_map[z["id_davky"]] = {"nazov": z.get("nazov_vyrobku"), "celkova": float(z.get("celkova_cena_surovin") or 0.0)}

    enriched_dmg = []
    for d in dmg_rows:
        zv_info = zv_map.get(d.get("id_davky"))
        nm  = (zv_info["nazov"] if zv_info and zv_info.get("nazov") else (d.get("nazov_vyrobku") or d.get("nazov_suroviny")))
        key = _norm_key(nm)
        cat = name2info.get(key, {}).get("kategoria", "Nezaradené")
        if use_cat and cat.casefold() != cat_norm:
            continue

        dd = dict(d)
        dd["kategoria_pre_recepty"] = cat
        dd["naklady_skody"]         = float((zv_info or {}).get("celkova") or d.get("naklady") or d.get("naklad") or d.get("cena") or 0.0)
        enriched_dmg.append(dd)

    return {"data": enriched_prod, "damage_data": enriched_dmg}

def get_receipt_report_html(period: str, category: str):
    """Tlačiteľný report príjmov do výrobného skladu (zaznamy_prijem)."""
    today = datetime.now().date()
    p = (period or '').strip().lower()
    if p in ('day','den','dnes','today'):
        start = end = today; label = today.strftime('%d.%m.%Y')
    elif p in ('week','tyzden','týždeň'):
        start = today - timedelta(days=today.weekday()); end = start + timedelta(days=6); label = f"{start.strftime('%d.%m.%Y')} – {end.strftime('%d.%m.%Y')}"
    elif p in ('month','mesiac'):
        start = today.replace(day=1); nxt = (start.replace(day=28)+timedelta(days=4)).replace(day=1); end = nxt - timedelta(days=1); label = start.strftime('%m/%Y')
    else:
        start = end = today; label = today.strftime('%d.%m.%Y')

    rows = db_connector.execute_query("""
        SELECT DATE(zp.datum) AS d, TIME(zp.datum) AS t, zp.nazov_suroviny AS name,
               COALESCE(zp.mnozstvo_kg,0) AS qty_kg, zp.nakupna_cena_eur_kg AS unit_price,
               zp.typ AS source, zp.poznamka_dodavatel AS note
        FROM zaznamy_prijem zp
        WHERE DATE(zp.datum) BETWEEN %s AND %s
        ORDER BY zp.datum ASC, zp.nazov_suroviny ASC
    """, (start, end)) or []

    total_qty = 0.0; total_val = 0.0
    trs=[]
    for r in rows:
        q=float(r.get('qty_kg') or 0.0); c=r.get('unit_price'); v=(q*float(c)) if c not in (None,'') else None
        total_qty+=q; total_val+=(v or 0.0)
        trs.append(
            f"<tr><td>{r['d'].strftime('%d.%m.%Y')}</td><td>{(r['t'].strftime('%H:%M') if r.get('t') else '')}</td>"
            f"<td>{html.escape(r['name'] or '')}</td>"
            f"<td style='text-align:right'>{q:.3f}</td>"
            f"<td style='text-align:right'>{'' if c in (None,'') else f'{float(c):.4f}'}</td>"
            f"<td style='text-align:right'>{'' if v is None else f'{v:.2f}'}</td>"
            f"<td>{html.escape(r.get('source') or '')}</td>"
            f"<td>{html.escape(r.get('note') or '')}</td></tr>"
        )

    html_out=f"""<!doctype html><html><head><meta charset="utf-8"><title>Príjem – report</title>
<style>body{{font-family:Inter,Arial,sans-serif;padding:16px}}table{{border-collapse:collapse;width:100%}}
th,td{{border:1px solid #e5e7eb;padding:6px 8px;text-align:left}}th{{background:#f9fafb}}</style></head><body>
<h2>Príjem do výrobného skladu</h2><p>Obdobie: {label}</p>
<table><thead><tr><th>Dátum</th><th>Čas</th><th>Surovina</th><th>Množstvo (kg)</th><th>Cena €/kg</th><th>Hodnota €</th><th>Zdroj</th><th>Poznámka</th></tr></thead>
<tbody>{''.join(trs) or "<tr><td colspan='8'>Žiadne príjmy.</td></tr>"}</tbody>
<tfoot><tr><th colspan="3" style="text-align:right">Súčet:</th><th style="text-align:right">{total_qty:.3f}</th><th></th><th style="text-align:right">{total_val:.2f}</th><th colspan="2"></th></tr></tfoot>
</table><script>window.print()</script></body></html>"""
    return make_response(html_out)

def get_inventory_difference_report_html(date_str: Optional[str]):
    """Tlačiteľný report inventúr Sklad 2 (expedicia_inventury + expedicia_inventura_polozky) za deň."""
    try:
        d = datetime.strptime((date_str or ''), "%Y-%m-%d").date() if date_str else datetime.now().date()
    except Exception:
        d = datetime.now().date()

    invs = db_connector.execute_query("""
        SELECT id, datum, vytvoril, created_at
        FROM expedicia_inventury
        WHERE datum = %s
        ORDER BY created_at ASC, id ASC
    """, (d,)) or []

    sections=[]; tot_abs=0.0; tot_val=0.0
    for inv in invs:
        lines = db_connector.execute_query("""
            SELECT ean, nazov, kategoria, system_stav_kg, realny_stav_kg, rozdiel_kg, hodnota_eur
            FROM expedicia_inventura_polozky
            WHERE inventura_id = %s
            ORDER BY kategoria, nazov
        """, (inv['id'],)) or []

        s_abs=sum(abs(float(r['rozdiel_kg'])) for r in lines); s_val=sum(float(r['hodnota_eur']) for r in lines)
        tot_abs+=s_abs; tot_val+=s_val
        trs="".join(
            f"<tr><td>{html.escape(r.get('ean') or '')}</td><td>{html.escape(r.get('nazov') or '')}</td><td>{html.escape(r.get('kategoria') or '')}</td>"
            f"<td style='text-align:right'>{float(r['system_stav_kg']):.3f}</td><td style='text-align:right'>{float(r['realny_stav_kg']):.3f}</td>"
            f"<td style='text-align:right'>{float(r['rozdiel_kg']):.3f}</td><td style='text-align:right'>{float(r['hodnota_eur']):.2f} €</td></tr>"
            for r in lines
        )
        created_s = inv.get('created_at').strftime('%H:%M') if isinstance(inv.get('created_at'), datetime) else ''
        sections.append(f"<h3>Inventúra ID {inv['id']} – {d.strftime('%d.%m.%Y')} (vytvoril: {html.escape(inv.get('vytvoril') or '')}, {created_s})</h3>"
                        f"<table><thead><tr><th>EAN</th><th>Názov</th><th>Kategória</th><th>Sys (kg)</th><th>Real (kg)</th><th>Rozdiel (kg)</th><th>Hodnota (€)</th></tr></thead>"
                        f"<tbody>{trs or '<tr><td colspan=7>Žiadne rozdiely.</td></tr>'}</tbody></table>")

    if not sections:
        sections.append("<p>V tento deň neexistuje žiadna inventúra.</p>")
    html_out=f"<!doctype html><html><head><meta charset='utf-8'><title>Inventúra – {d.strftime('%d.%m.%Y')}</title></head><body>{''.join(sections)}<script>window.print()</script></body></html>"
    return make_response(html_out)


# =================================================================
# === PRÍJEM Z VÝROBY – report podľa dátumu PRÍJMU (Expedícia) =====
# =================================================================

def _oh_has_col(table: str, col: str) -> bool:
    return _has_col(table, col)

def _oh_zv_name_col() -> str:
    return 'nazov_vyrobu' if _oh_has_col('zaznamy_vyroba','nazov_vyrobu') else 'nazov_vyrobku'

def _oh_safe_float(x) -> float:
    try: return float(x)
    except Exception: return 0.0

def get_reception_report(date_from: str, date_to: str, overhead_coeff: float = 1.15):
    """
    Príjem z výroby podľa dátumu PRÍJMU (Expedícia) – vážené ceny aj bez uzávierky dňa.
    Jednotková cena:
      1) zv.cena_za_jednotku (>0), inak
      2) zv.celkova_cena_surovin / reálne prijaté (kg/ks)
    """
    if not date_from or not date_to:
        return {"error": "Zadajte date_from a date_to (YYYY-MM-DD)."}
    try:
        d_from = datetime.strptime(date_from, "%Y-%m-%d").strftime("%Y-%m-%d")
        d_to   = datetime.strptime(date_to,   "%Y-%m-%d").strftime("%Y-%m-%d")
    except Exception:
        return {"error": "Neplatný formát dátumu. Použite YYYY-MM-DD."}

    zv = _oh_zv_name_col()

    rows = db_connector.execute_query(f"""
        SELECT
            ep.id_davky, ep.unit AS ep_unit, ep.prijem_kg, ep.prijem_ks, ep.datum_prijmu,
            zv.{zv} AS product_name, zv.planovane_mnozstvo_kg AS planned_kg_batch,
            COALESCE(zv.celkova_cena_surovin,0) AS zv_total_cost,
            COALESCE(zv.cena_za_jednotku,0)     AS zv_unit_cost,
            COALESCE(zv.realne_mnozstvo_kg,0)   AS zv_real_kg,
            COALESCE(zv.realne_mnozstvo_ks,0)   AS zv_real_ks,
            p.mj AS prod_mj, COALESCE(p.vaha_balenia_g,0) AS weight_g
        FROM expedicia_prijmy ep
        JOIN zaznamy_vyroba zv ON zv.id_davky = ep.id_davky
        JOIN produkty      p  ON TRIM(zv.{zv}) = TRIM(p.nazov_vyrobku)
        WHERE ep.is_deleted = 0 AND ep.datum_prijmu BETWEEN %s AND %s
    """, (d_from, d_to)) or []

    if not rows:
        return {"period":{"date_from":d_from,"date_to":d_to,"overhead_coeff":overhead_coeff},
                "rows":[], "totals":{"planned_kg":0.0,"real_kg":0.0,"yield_pct":0.0}}

    totals_per_batch = db_connector.execute_query(f"""
        SELECT ep.id_davky,
               SUM(CASE WHEN ep.unit='kg' THEN COALESCE(ep.prijem_kg,0)
                        ELSE COALESCE(ep.prijem_ks,0) * COALESCE(p.vaha_balenia_g,0) / 1000 END) AS total_real_kg,
               SUM(CASE WHEN ep.unit='ks' THEN COALESCE(ep.prijem_ks,0) ELSE 0 END) AS total_real_ks
          FROM expedicia_prijmy ep
          JOIN zaznamy_vyroba zv ON zv.id_davky = ep.id_davky
          JOIN produkty p        ON TRIM(zv.{zv}) = TRIM(p.nazov_vyrobku)
         WHERE ep.is_deleted = 0
         GROUP BY ep.id_davky
    """) or []
    total_map = {r['id_davky']: (_oh_safe_float(r['total_real_kg']), _oh_safe_float(r['total_real_ks'])) for r in totals_per_batch}

    agg: Dict[tuple, Dict[str, float]] = {}
    total_planned, total_real = 0.0, 0.0

    for r in rows:
        product = str(r['product_name']).strip()
        mj      = r['prod_mj'] or 'kg'
        w_g     = _oh_safe_float(r['weight_g'])
        planned_batch = _oh_safe_float(r['planned_kg_batch'])
        tot_kg, tot_ks = total_map.get(r['id_davky'], (0.0, 0.0))

        real_kg = _oh_safe_float(r['prijem_kg']) if r['ep_unit']=='kg' else (_oh_safe_float(r['prijem_ks']) * w_g / 1000.0)
        planned_prop = planned_batch * (real_kg / tot_kg) if tot_kg > 0 else real_kg

        unit_cost_for_batch = _oh_safe_float(r['zv_unit_cost'])
        if unit_cost_for_batch <= 0.0 and _oh_safe_float(r['zv_total_cost']) > 0.0:
            if mj == 'kg':
                denom = _oh_safe_float(r['zv_real_kg']) if _oh_safe_float(r['zv_real_kg'])>0 else tot_kg
            else:
                denom = _oh_safe_float(r['zv_real_ks']) if _oh_safe_float(r['zv_real_ks'])>0 else tot_ks
            if denom > 0:
                unit_cost_for_batch = _oh_safe_float(r['zv_total_cost']) / denom

        if mj == 'kg':
            units_for_avg = real_kg
        else:
            units_for_avg = _oh_safe_float(r['prijem_ks']) if r['ep_unit']=='ks' else ((real_kg / (w_g/1000.0)) if w_g > 0 else 0.0)

        key = (product, mj)
        rec = agg.setdefault(key, {"planned_kg":0.0,"real_kg":0.0,"sum_cost_x_units":0.0,"sum_units":0.0})
        rec["planned_kg"] += planned_prop
        rec["real_kg"]    += real_kg
        if unit_cost_for_batch > 0 and units_for_avg > 0:
            rec["sum_cost_x_units"] += unit_cost_for_batch * units_for_avg
            rec["sum_units"]        += units_for_avg

    out_rows = []
    for (product, mj), s in sorted(agg.items(), key=lambda x: x[0][0].lower()):
        planned = s["planned_kg"]; real = s["real_kg"]
        total_planned += planned; total_real += real
        unit_cost_no_overhead = (s["sum_cost_x_units"] / s["sum_units"]) if s["sum_units"] > 0 else 0.0
        out_rows.append({
            "product": product, "unit": mj,
            "planned_kg": round(planned,3), "real_kg": round(real,3),
            "yield_pct": round((real/planned*100.0),2) if planned>0 else None,
            "unit_cost_no_overhead": round(unit_cost_no_overhead,4),
            "unit_cost_with_overhead": round(unit_cost_no_overhead * float(overhead_coeff or 1.0),4),
        })

    total_yield = (total_real/total_planned*100.0) if total_planned>0 else 0.0
    return {
        "period": {"date_from": d_from, "date_to": d_to, "overhead_coeff": float(overhead_coeff or 1.0)},
        "rows": out_rows,
        "totals": {"planned_kg":round(total_planned,3),"real_kg":round(total_real,3),"yield_pct":round(total_yield,2)}
    }


# =================================================================
# === B2C PRICELIST / OBJEDNÁVKY / ODMENY ==========================
# =================================================================

def get_b2c_pricelist_for_admin():
    all_products = db_connector.execute_query("""
        SELECT ean, nazov_vyrobku, predajna_kategoria, dph, mj
        FROM produkty
        WHERE typ_polozky='produkt' OR typ_polozky LIKE 'VÝROBOK%%' OR typ_polozky LIKE 'TOVAR%%'
        ORDER BY predajna_kategoria, nazov_vyrobku
    """) or []
    items = db_connector.execute_query(f"""
        SELECT c.ean_produktu, p.nazov_vyrobku, p.dph,
               c.cena_bez_dph, c.je_v_akcii, c.akciova_cena_bez_dph
          FROM b2c_cennik_polozky c
          JOIN produkty p ON c.ean_produktu COLLATE {COLL} = p.ean COLLATE {COLL}
    """) or []

    avg_idx = _avg_costs_index_by_ean()
    for p in all_products:
        rec = avg_idx.get((p.get("ean") or "").strip())
        p["default_unit_cost"] = (rec or {}).get("avg_manufacturing_unit_cost")

    return {"all_products": all_products, "pricelist": items}

def update_b2c_pricelist(data):
    items = (data or {}).get('items') or []
    if not items:
        return {"error": "Nie je čo uložiť."}
    meta = _b2c_meta_load()
    for it in items:
        ean = (it.get('ean') or '').strip()
        if not ean:
            continue
        price      = it.get('price')
        is_akcia   = 1 if it.get('is_akcia') else 0
        sale_price = it.get('sale_price') or None
        desc       = (it.get('description') or '').strip()
        img_url    = (it.get('image_url') or '').strip()

        db_connector.execute_query("""
            UPDATE b2c_cennik_polozky
               SET cena_bez_dph=%s, je_v_akcii=%s, akciova_cena_bez_dph=%s
             WHERE ean_produktu=%s
        """, (price, is_akcia, sale_price, ean), fetch='none')

        meta.setdefault(ean, {})
        if desc != '':
            meta[ean]['popis'] = desc
        if img_url != '':
            meta[ean]['obrazok'] = img_url

    _b2c_meta_save(meta)
    return {"message": "Zmeny v cenníku uložené."}

def add_products_to_b2c_pricelist(data):
    items = (data or {}).get('items') or []
    if not items:
        return {"error": "Žiadne položky na pridanie."}
    inserted = 0; updated  = 0
    for it in items:
        ean = str(it.get('ean') or '').strip()
        if not ean:
            continue
        price = it.get('price')
        try:
            price = float(price) if price is not None and str(price).strip() != '' else 0.0
        except Exception:
            price = 0.0

        exists = db_connector.execute_query(
            f"SELECT 1 FROM b2c_cennik_polozky WHERE ean_produktu COLLATE {COLL} = %s LIMIT 1",
            (ean,)
        )
        if exists:
            db_connector.execute_query(
                "UPDATE b2c_cennik_polozky SET cena_bez_dph = %s WHERE ean_produktu = %s",
                (price, ean), fetch='none'
            )
            updated += 1
        else:
            db_connector.execute_query(
                "INSERT INTO b2c_cennik_polozky (ean_produktu, cena_bez_dph, je_v_akcii, akciova_cena_bez_dph) VALUES (%s, %s, 0, NULL)",
                (ean, price),
                fetch='none'
            )
            inserted += 1

    return {"message": "Hotovo.", "inserted": inserted, "updated": updated}

def get_b2c_orders_for_admin():
    return db_connector.execute_query("""
        SELECT o.*, z.nazov_firmy as zakaznik_meno
        FROM b2c_objednavky o
        JOIN b2b_zakaznici z ON o.zakaznik_id = z.zakaznik_id
        ORDER BY 
            CASE o.stav
                WHEN 'Prijatá' THEN 1
                WHEN 'Pripravená' THEN 2
                WHEN 'Hotová' THEN 3
                WHEN 'Zrušená' THEN 4
                ELSE 5
            END,
            o.pozadovany_datum_dodania ASC,
            o.datum_objednavky ASC
    """) or []

def update_b2c_order_status(data):
    """
    Zmení stav B2C objednávky. NOTIFIKÁCIE SA TU NEPOSIELAJÚ,
    aby READY nechodilo 2× (odosiela sa výhradne cez /email/ready a /sms/ready).
    """
    order_id = (data or {}).get('order_id')
    new_status = (data or {}).get('status')
    if not order_id or not new_status:
        return {"error": "Chýba ID alebo nový stav."}

    new_status = str(new_status)
    if new_status not in ('Prijatá', 'Pripravená', 'Hotová', 'Zrušená'):
        return {"error": "Neplatný stav."}

    try:
        db_connector.execute_query(
            "UPDATE b2c_objednavky SET stav=%s WHERE id=%s", (new_status, order_id), fetch='none'
        )
    except Exception as e:
        return {"error": f"Nepodarilo sa zapísať stav: {e}"}

    return {"message": f"Stav objednávky zmenený na '{new_status}'."}

def finalize_b2c_order(data):
    """
    Nastaví objednávku do stavu 'Pripravená' a uloží finálnu cenu.
    NOTIFIKÁCIE SA TU NEPOSIELAJÚ (aby neboli duplicity).
    Vstup: { "order_id": int, "final_price": "12.34" }
    """
    order_id = (data or {}).get('order_id')
    price_raw = (data or {}).get('final_price')
    if not order_id or price_raw in (None, ''):
        return {"error": "Chýba ID objednávky alebo finálna cena."}

    try:
        final_price = float(str(price_raw).replace(",", "."))
        if final_price <= 0:
            return {"error": "Finálna cena musí byť kladné číslo."}
    except Exception:
        return {"error": "Neplatný formát finálnej ceny."}

    # Ulož stav
    db_connector.execute_query(
        "UPDATE b2c_objednavky SET stav=%s WHERE id=%s", ('Pripravená', order_id), fetch='none'
    )

    # Ulož finálnu cenu do prvého dostupného stĺpca (žiadne ALTERy)
    for col in ("celkova_suma_s_dph", "finalna_suma_s_dph", "suma_s_dph", "total_s_dph", "finalna_suma", "total_gross"):
        try:
            db_connector.execute_query(
                f"UPDATE b2c_objednavky SET {col}=%s WHERE id=%s",
                (final_price, order_id), fetch='none'
            )
            break
        except Exception:
            continue

    return {"message": "Objednávka je označená ako 'Pripravená' a cena uložená.", "final_price": final_price}


def credit_b2c_loyalty_points(data):
    """
    Pripíše vernostné body podľa finálnej ceny a nastaví stav 'Hotová'.
    Po úspechu odošle B2C COMPLETED SMS (uvedie uhradené € a body).
    Vstup: { "order_id": int }
    """
    order_id = data.get('order_id')
    if not order_id:
        return {"error": "Chýba ID objednávky."}

    order = db_connector.execute_query("SELECT * FROM b2c_objednavky WHERE id = %s", (order_id,), 'one')
    if not order:
        return {"error": "Objednávka nebola nájdená."}
    if order.get('stav') not in ('Pripravená', 'Hotová'):
        return {"error": "Body je možné pripísať len pre 'Pripravená'."}
    if order.get('stav') == 'Hotová':
        return {"message": "Objednávka je už v stave 'Hotová'."}

    try:
        final_price = float(order.get('celkova_suma_s_dph') or 0.0)
    except Exception:
        final_price = 0.0
    if final_price <= 0:
        return {"error": "Objednávka nemá evidovanú celkovú cenu."}

    points_to_add = math.floor(final_price)
    cust_login = order.get('zakaznik_id')

    db_connector.execute_query(
        "UPDATE b2b_zakaznici SET vernostne_body = COALESCE(vernostne_body,0) + %s WHERE zakaznik_id = %s",
        (points_to_add, cust_login), fetch='none'
    )
    db_connector.execute_query(
        "UPDATE b2c_objednavky SET stav = 'Hotová' WHERE id = %s",
        (order_id,), fetch='none'
    )

    # SERVER-SIDE SMS AUTONOTIFY (COMPLETED)
    try:
        info = db_connector.execute_query(
            "SELECT cislo_objednavky FROM b2c_objednavky WHERE id=%s",
            (order_id,), fetch="one"
        ) or {}
        notify = globals().get('b2c_notify_sms_completed')
        if callable(notify):
            notify({
                "order_id": order_id,
                "order_no": info.get("cislo_objednavky"),
                "final_paid": final_price,
                "points_added": points_to_add,
                "customer_ref": cust_login,
            })
    except Exception as _e:
        print("b2c_notify_sms_completed error:", _e)

    return {"message": f"Pripísaných {points_to_add} bodov. Objednávka je 'Hotová'."}

def b2c_notify_sms_cancelled(data: dict):
    """
    Posle SMS o zruseni objednavky s dovodom (bez diakritiky).
    Vstup: { "order_id": int, "order_no": str?, "reason": str }
    """
    od = data or {}
    info = _sms_get_b2c_order_info(od.get("order_id"), od.get("order_no"))
    order_no = info.get("order_no") or (od.get("order_no") or "")
    if not order_no:
        return {"error": "Chyba: nepozname cislo objednavky."}

    reason = (od.get("reason") or od.get("dovod") or od.get("dovod_neprevzatia") or "").strip()
    if not reason:
        reason = "neprevziatie"

    # idempotencia
    meta = _sms_read_order_meta(order_no)
    mp = meta.get("_meta_path")
    if meta.get("cancelled_sms_sent_at"):
        return {"id":"SKIPPED_DUP","note":"CANCELLED SMS uz odoslana","order_no":order_no}

    phone = _sms_resolve_b2c_phone(None, order_no, od.get("phone"), od.get("user_email"), info.get("customer_ref"))
    if not phone:
        return {"error":"Cislo zakaznika sa nenaslo","order_no":order_no}

    # skratime reason do cca 90 znakov
    reason_s = reason.replace("\n"," ").replace("\r"," ")
    if len(reason_s) > 90:
        reason_s = reason_s[:87] + "..."

    txt = f"MIK: Objednavka {order_no} bola zrusena. Dovod: {reason_s}"
    lock = _sms_try_lock(order_no, "cancelled", ttl_sec=30)
    if not lock:
        return {"id":"SKIPPED_LOCK","note":"CANCELLED dup <30s","order_no":order_no}
    try:
        res = _sms_send(txt, phone)
        try:
            meta["cancelled_sms_sent_at"] = datetime.utcnow().isoformat() + "Z"
            with open(mp, "w", encoding="utf-8") as f:
                json.dump(meta, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        return res
    finally:
        _sms_unlock(lock)


def cancel_b2c_order(data):
    order_id, reason = data.get('order_id'), data.get('reason')
    if not all([order_id, reason]): return {"error": "Chýba ID objednávky alebo dôvod zrušenia."}
    db_connector.execute_query(
        "UPDATE b2c_objednavky SET stav = 'Zrušená', poznamka = CONCAT(IFNULL(poznamka, ''), ' | ZRUŠENÉ: ', %s) WHERE id = %s",
        (reason, order_id), fetch='none'
    )
    order = db_connector.execute_query("SELECT zakaznik_id, cislo_objednavky FROM b2c_objednavky WHERE id = %s", (order_id,), 'one')
    if order:
        customer = db_connector.execute_query("SELECT nazov_firmy, email FROM b2b_zakaznici WHERE zakaznik_id = %s", (order['zakaznik_id'],), 'one')
        if customer:
            try:
                notification_handler.send_b2c_order_cancelled_email(customer['email'], customer['nazov_firmy'], order['cislo_objednavky'], reason)
            except Exception as e:
                print(f"B2C cancel email err: {e}")
    return {"message": "Objednávka bola zrušená a zákazník notifikovaný."}

def get_b2c_customers_for_admin():
    return db_connector.execute_query(
        "SELECT zakaznik_id, nazov_firmy, email, telefon, adresa, adresa_dorucenia, vernostne_body FROM b2b_zakaznici WHERE typ = 'B2C' ORDER BY nazov_firmy"
    ) or []


# =================================================================
# === B2C – REWARDS (schema + CRUD) ================================
# =================================================================

def _ensure_b2c_rewards_schema():
    db_connector.execute_query("""
        CREATE TABLE IF NOT EXISTS b2c_vernostne_odmeny (
          id INT AUTO_INCREMENT PRIMARY KEY,
          nazov_odmeny VARCHAR(255) NOT NULL,
          potrebne_body INT NOT NULL,
          je_aktivna TINYINT(1) NOT NULL DEFAULT 1,
          valid_from DATE NULL,
          valid_to DATE NULL,
          created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
          updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_slovak_ci
    """, fetch='none')

def _has_col_b2c(table, col):
    return _has_col(table, col)

def get_b2c_rewards_for_admin():
    _ensure_b2c_rewards_schema()
    has_from = _has_col_b2c("b2c_vernostne_odmeny", "valid_from")
    has_to   = _has_col_b2c("b2c_vernostne_odmeny", "valid_to")
    cols = [ "id", "nazov_odmeny", "potrebne_body", "je_aktivna",
             ("valid_from" if has_from else "NULL AS valid_from"),
             ("valid_to"   if has_to   else "NULL AS valid_to") ]
    sql = f"SELECT {', '.join(cols)} FROM b2c_vernostne_odmeny ORDER BY id DESC"
    return db_connector.execute_query(sql) or []

def add_b2c_reward(data):
    _ensure_b2c_rewards_schema()
    name   = (data or {}).get('name', '').strip()
    points = (data or {}).get('points', '')
    if not name or str(points).strip()=='':
        return {"error": "Chýba názov odmeny alebo počet bodov."}
    try:
        pts = int(float(str(points).replace(',', '.')))
        if pts <= 0: raise ValueError
    except Exception:
        return {"error": "Počet bodov musí byť kladné číslo."}

    has_from = _has_col_b2c("b2c_vernostne_odmeny", "valid_from")
    has_to   = _has_col_b2c("b2c_vernostne_odmeny", "valid_to")

    fields = ["nazov_odmeny", "potrebne_body", "je_aktivna"]
    vals   = [name, pts, 1]; ph = ["%s","%s","%s"]
    if has_from: fields.append("valid_from"); vals.append((data.get('valid_from') or None)); ph.append("%s")
    if has_to:   fields.append("valid_to");   vals.append((data.get('valid_to') or None));   ph.append("%s")

    sql = f"INSERT INTO b2c_vernostne_odmeny ({', '.join(fields)}) VALUES ({', '.join(ph)})"
    db_connector.execute_query(sql, tuple(vals), fetch='none')
    return {"message": "Odmena bola vytvorená."}

def toggle_b2c_reward_status(data):
    _ensure_b2c_rewards_schema()
    rid    = (data or {}).get('id')
    status = (data or {}).get('status')
    if rid is None:
        return {"error": "Chýba ID odmeny."}
    s = str(status).strip().lower()
    cur_on = (status is True) or (s in ('1','true','t','yes','y','on'))
    new_status = 0 if cur_on else 1
    db_connector.execute_query("UPDATE b2c_vernostne_odmeny SET je_aktivna=%s WHERE id=%s", (new_status, rid), fetch='none')
    return {"message": "Stav odmeny bol zmenený.", "new_status": new_status}

def update_b2c_reward(data):
    _ensure_b2c_rewards_schema()
    rid = (data or {}).get('id')
    if not rid:
        return {"error": "Chýba ID odmeny."}
    name = (data or {}).get('name','').strip()
    points = (data or {}).get('points','')
    if not name:
        return {"error":"Názov odmeny je povinný."}
    try:
        pts = int(float(str(points).replace(',', '.')));  assert pts>0
    except Exception:
        return {"error":"Počet bodov musí byť kladné číslo."}

    has_from = _has_col_b2c("b2c_vernostne_odmeny", "valid_from")
    has_to   = _has_col_b2c("b2c_vernostne_odmeny", "valid_to")

    cols = ["nazov_odmeny=%s", "potrebne_body=%s"]; params = [name, pts]
    if has_from: cols.append("valid_from=%s"); params.append((data.get('valid_from') or None))
    if has_to:   cols.append("valid_to=%s");   params.append((data.get('valid_to') or None))
    params.append(rid)
    sql = f"UPDATE b2c_vernostne_odmeny SET {', '.join(cols)} WHERE id=%s"
    db_connector.execute_query(sql, tuple(params), fetch='none')
    return {"message": "Odmena upravená."}

def edit_b2c_reward(data):
    return update_b2c_reward(data)


# =================================================================
# === HACCP – JSON úložisko + DOCX import/export ===================
# =================================================================

def _haccp_dir():
    base = os.path.dirname(__file__)
    d = os.path.join(base, 'data', 'haccp')
    os.makedirs(d, exist_ok=True)
    return d

def _haccp_upload_dir():
    base = os.path.dirname(__file__)
    d = os.path.join(base, 'static', 'uploads', 'haccp')
    os.makedirs(d, exist_ok=True)
    return d

def _load_haccp_json(doc_id):
    path = os.path.join(_haccp_dir(), f"{doc_id}.json")
    if not os.path.exists(path):
        return None, path
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f), path

def _save_haccp_json(doc_obj, dst_path=None):
    if not dst_path:
        dst_path = os.path.join(_haccp_dir(), f"{doc_obj['id']}.json")
    with open(dst_path, 'w', encoding='utf-8') as f:
        json.dump(doc_obj, f, ensure_ascii=False, indent=2)
    return dst_path

def _now_iso():
    return datetime.utcnow().isoformat(timespec='seconds') + 'Z'

def _docx_to_html_best_effort(docx_path: str) -> str:
    """Konverzia .docx -> HTML (1) mammoth, (2) python-docx, (3) XML fallback)."""
    try:
        import mammoth
        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_html(f, convert_image=mammoth.images.inline())
        html_out = (result.value or "").strip()
        if html_out:
            return html_out
    except Exception:
        pass

    try:
        import docx as docx_lib
        doc = docx_lib.Document(docx_path)
        parts = []
        for p in doc.paragraphs:
            style = (getattr(p.style, 'name', '') or '').lower()
            tag = 'p'
            for n, t in (('1','h1'),('2','h2'),('3','h3'),('4','h4')):
                if f'heading {n}' in style or f'nadpis {n}' in style:
                    tag = t; break
            runs_html = []
            for r in p.runs:
                tt = html.escape(r.text or '').replace('\n','<br>')
                if not tt: continue
                if r.bold: tt = f"<strong>{tt}</strong>"
                if r.italic: tt = f"<em>{tt}</em>"
                if r.underline: tt = f"<u>{tt}</u>"
                runs_html.append(tt)
            parts.append(f"<{tag}>{''.join(runs_html) or '&nbsp;'}</{tag}>")
        html_out = ''.join(parts).strip()
        if html_out: return html_out
    except Exception:
        pass

    try:
        with zipfile.ZipFile(docx_path) as z:
            with z.open('word/document.xml') as fh:
                tree = ET.parse(fh)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        body = tree.getroot().find('.//w:body', ns)
        out = []
        if body is not None:
            for p in body.findall('w:p', ns):
                chunks = []
                for node in p.iter():
                    tag = node.tag
                    if tag == f"{{{ns['w']}}}t":
                        chunks.append(html.escape(node.text or ''))
                    elif tag == f"{{{ns['w']}}}br":
                        chunks.append('<br>')
                para = ''.join(chunks).strip()
                out.append(f"<p>{para or '&nbsp;'}</p>")
        html_out = ''.join(out).strip()
        if html_out: return html_out
    except Exception:
        pass

    return "<p>(prázdny dokument)</p>"

def import_haccp_docx():
    file = request.files.get('file')
    if not file or not file.filename:
        return {"error":"Chýba súbor .docx (pole 'file')."}, 400
    name = file.filename
    ext = os.path.splitext(name)[1].lower()
    if ext != '.docx':
        return {"error":"Podporovaný je len formát .docx."}, 400

    up_dir = _haccp_upload_dir()
    saved_name = f"{uuid.uuid4().hex}.docx"
    saved_path = os.path.join(up_dir, saved_name)
    file.save(saved_path)
    original_url = f"/static/uploads/haccp/{saved_name}"

    content_html = _docx_to_html_best_effort(saved_path)

    doc_id = uuid.uuid4().hex
    title = (request.form.get('title') or os.path.splitext(name)[0]).strip() or "HACCP dokument"
    now = _now_iso()
    obj = {
        "id": doc_id,
        "title": title,
        "content": content_html,
        "created_at": now,
        "updated_at": now,
        "attachments": {"original_docx": original_url}
    }
    dst = os.path.join(_haccp_dir(), f"{doc_id}.json")
    with open(dst, 'w', encoding='utf-8') as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)

    return {"message": "Dokument importovaný.", "doc": {"id": doc_id, "title": title}}

def _html_to_text(html_content: str) -> str:
    txt_content = html_content or ""
    txt_content = re.sub(r'(?i)<br\s*/?>', '\n', txt_content)
    txt_content = re.sub(r'(?i)</p\s*>', '\n', txt_content)
    txt_content = re.sub(r'<[^>]+>', '', txt_content)
    txt_content = html.unescape(txt_content)
    txt_content = txt_content.replace('\r\n', '\n').replace('\r', '\n')
    return txt_content.strip()

def _write_minimal_docx(text_content: str, out_path: str):
    def xml_escape(s: str) -> str:
        return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    paragraphs_xml = "".join(
        f"<w:p><w:r><w:t xml:space='preserve'>{xml_escape(line)}</w:t></w:r></w:p>"
        for line in (text_content.split("\n") if text_content else [""])
    )
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{paragraphs_xml}<w:sectPr/></w:body>
</w:document>"""
    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>"""
    rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    doc_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>"""
    styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:styles/>
</w:styles>"""
    with zipfile.ZipFile(out_path,"w",compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types_xml)
        z.writestr("_rels/.rels", rels_xml)
        z.writestr("word/_rels/document.xml.rels", doc_rels_xml)
        z.writestr("word/document.xml", document_xml)
        z.writestr("word/styles.xml", styles_xml)
    return out_path

def export_haccp_docx():
    """GET: id (povinné), use_original=1 (voliteľné) -> vráti .docx na stiahnutie."""
    doc_id = request.args.get('id') or request.values.get('id')
    if not doc_id:
        return {"error":"Chýba parameter 'id'."}, 400

    obj, _ = _load_haccp_json(doc_id)
    if not obj:
        return {"error":"Dokument neexistuje."}, 404

    if request.args.get('use_original') == '1':
        url = ((obj.get("attachments") or {}).get("original_docx"))
        if url:
            base = os.path.dirname(__file__)
            fs_path = os.path.normpath(os.path.join(base, url.lstrip('/')))
            if os.path.exists(fs_path):
                return send_file(fs_path, as_attachment=True, download_name=f"{obj.get('title','haccp')}.docx")

    text = _html_to_text(obj.get("content") or "")
    tmp_dir = os.path.join(_haccp_dir(), "_export"); os.makedirs(tmp_dir, exist_ok=True)
    out_path = os.path.join(tmp_dir, f"{doc_id}.docx")

    try:
        import docx
        d = docx.Document()
        for line in (text.split("\n") if text else [""]):
            d.add_paragraph(line)
        d.save(out_path)
    except Exception:
        _write_minimal_docx(text, out_path)

    return send_file(out_path, as_attachment=True, download_name=f"{obj.get('title','haccp')}.docx")

# ---- HACCP JSON API pre UI (get/list/save) -----------------------

def get_haccp_docs():
    """Zoznam HACCP JSON dokumentov (id, title, created/updated)."""
    docs = []
    base = _haccp_dir()
    try:
        for fn in os.listdir(base):
            if not fn.endswith('.json'): continue
            path = os.path.join(base, fn)
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    obj = json.load(f)
                docs.append({
                    "id": obj.get("id") or os.path.splitext(fn)[0],
                    "title": obj.get("title") or os.path.splitext(fn)[0],
                    "created_at": obj.get("created_at"),
                    "updated_at": obj.get("updated_at"),
                })
            except Exception:
                continue
    except Exception:
        pass
    docs.sort(key=lambda d: (d.get("updated_at") or d.get("created_at") or ""), reverse=True)
    return {"docs": docs}

def get_haccp_doc_content(payload: Dict[str, Any]):
    doc_id = (payload or {}).get('id')
    if not doc_id: return {"error":"Chýba id."}
    obj, _ = _load_haccp_json(doc_id)
    if not obj: return {"error":"Dokument neexistuje."}
    return {"doc": obj}

def save_haccp_doc(payload: Dict[str, Any]):
    """Uloží obsah HACCP JSON dokumentu. Očakáva: {id, title?, content?}"""
    doc_id = (payload or {}).get('id')
    if not doc_id: return {"error":"Chýba id."}
    obj, path = _load_haccp_json(doc_id)
    if not obj:
        obj = {"id": doc_id, "title": (payload or {}).get('title') or "HACCP dokument",
               "content": (payload or {}).get('content') or "", "created_at": _now_iso(), "updated_at": _now_iso()}
    else:
        if 'title' in payload:   obj['title'] = (payload.get('title') or obj.get('title') or "")
        if 'content' in payload: obj['content'] = (payload.get('content') or obj.get('content') or "")
        obj['updated_at'] = _now_iso()
    _save_haccp_json(obj, path)
    return {"message":"Uložené.", "doc": obj}


# =================================================================
# === DODÁVATELIA – SUPPLIERS (CRUD) ===============================
# =================================================================

def _ensure_suppliers_schema():
    db_connector.execute_query("""
        CREATE TABLE IF NOT EXISTS suppliers (
          id INT AUTO_INCREMENT PRIMARY KEY,
          name VARCHAR(255) NOT NULL,
          phone VARCHAR(64),
          email VARCHAR(255),
          address VARCHAR(255),
          is_active TINYINT(1) NOT NULL DEFAULT 1,
          created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
          updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_slovak_ci
    """, fetch='none')
    db_connector.execute_query("""
        CREATE TABLE IF NOT EXISTS supplier_categories (
          supplier_id INT NOT NULL,
          category VARCHAR(64) NOT NULL,
          PRIMARY KEY (supplier_id, category),
          CONSTRAINT fk_supcat_supplier FOREIGN KEY (supplier_id)
            REFERENCES suppliers(id) ON DELETE CASCADE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_slovak_ci
    """, fetch='none')

def list_suppliers(category: Optional[str] = None) -> Dict[str, Any]:
    _ensure_suppliers_schema()
    rows = db_connector.execute_query("SELECT id, name, phone, email, address FROM suppliers WHERE is_active=1 ORDER BY name") or []
    cats = db_connector.execute_query("SELECT supplier_id, category FROM supplier_categories") or []
    by: Dict[int, List[str]] = {}
    for c in cats:
        by.setdefault(c['supplier_id'], []).append(c['category'])
    cat = (category or '').strip().lower() or None
    items = []
    for r in rows:
        cset = by.get(r['id'], [])
        if cat and cat not in cset:
            continue
        items.append({
            "id": r["id"], "name": r["name"],
            "phone": r.get("phone"), "email": r.get("email"), "address": r.get("address"),
            "categories": cset
        })
    return {"items": items}

def create_supplier(payload: Dict[str, Any]) -> Dict[str, Any]:
    _ensure_suppliers_schema()
    name = (payload or {}).get('name', '').strip()
    if not name:
        return {"error": "Názov dodávateľa je povinný."}
    phone       = (payload or {}).get('phone')
    email_val   = (payload or {}).get('email')
    address_val = (payload or {}).get('address')
    categories  = (payload or {}).get('categories') or []

    new_id = db_connector.execute_query(
        "INSERT INTO suppliers (name, phone, email, address, is_active, created_at, updated_at) VALUES (%s,%s,%s,%s,1,NOW(),NOW())",
        (name, phone, email_val, address_val), fetch='lastrowid'
    )
    if categories:
        db_connector.execute_query(
            "INSERT INTO supplier_categories (supplier_id, category) VALUES (%s,%s)",
            [(new_id, c) for c in categories], multi=True, fetch='none'
        )
    return {"message": "Dodávateľ pridaný.", "id": new_id}

def update_supplier(payload: Dict[str, Any]) -> Dict[str, Any]:
    _ensure_suppliers_schema()

    sup_id = (payload or {}).get("id")
    if not sup_id:
        return {"error": "Chýba ID dodávateľa."}

    name = (payload or {}).get("name", "").strip()
    if not name:
        return {"error": "Názov dodávateľa je povinný."}

    phone        = (payload or {}).get("phone")
    email_val    = (payload or {}).get("email")
    address_val  = (payload or {}).get("address")
    categories   = (payload or {}).get("categories") or []

    db_connector.execute_query(
        "UPDATE suppliers SET name=%s, phone=%s, email=%s, address=%s, updated_at=NOW() WHERE id=%s",
        (name, phone, email_val, address_val, sup_id),
        fetch='none'
    )

    db_connector.execute_query("DELETE FROM supplier_categories WHERE supplier_id=%s", (sup_id,), fetch='none')
    if categories:
        db_connector.execute_query(
            "INSERT INTO supplier_categories (supplier_id, category) VALUES (%s,%s)",
            [(sup_id, c) for c in categories],
            multi=True,
            fetch='none'
        )

    return {"message": "Dodávateľ upravený."}

def delete_supplier(supplier_id: int) -> Dict[str, Any]:
    _ensure_suppliers_schema()
    db_connector.execute_query("UPDATE suppliers SET is_active=0, updated_at=NOW() WHERE id=%s", (supplier_id,), fetch='none')
    return {"message": "Dodávateľ zmazaný."}
# ============================================================================
# ======================  B2C SMS AUTO-NOTIFY  ===============================
# ============================================================================
# Tento blok rieši len SMS notifikácie pre B2C. Ostatné funkcie v module nemeň.

from typing import Optional
import time

# ---- lazy import smstools klienta (sms_handler.py) -------------------------
try:
    import sms_handler as _sms_mod
except Exception:
    _sms_mod = None

# ---- Low-level util --------------------------------------------------------
def _sms_table_exists(name: str) -> bool:
    r = db_connector.execute_query(
        "SELECT COUNT(*) c FROM information_schema.TABLES WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME=%s",
        (name,), fetch="one"
    )
    return bool(r and int(list(r.values())[0]) > 0)

def _sms_cols_for(name: str) -> set:
    rows = db_connector.execute_query(
        "SELECT COLUMN_NAME FROM information_schema.COLUMNS WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME=%s",
        (name,), fetch="all"
    ) or []
    return {r["COLUMN_NAME"] for r in rows}

def _sms_pick(colset: set, *cands, default=None):
    for c in cands:
        if c in colset:
            return c
    return default

def _sms_normalize_msisdn(number: Optional[str]) -> Optional[str]:
    try:
        return _sms_mod.normalize_msisdn(number) if (_sms_mod and number) else None
    except Exception:
        return None

def _sms_read_order_meta(order_no: Optional[str]) -> dict:
    if not order_no: return {}
    safe = "".join(ch for ch in str(order_no) if ch.isalnum() or ch in ("-","_"))
    path = os.path.join(os.path.dirname(__file__), "static", "uploads", "orders", f"{safe}.meta.json")
    try:
        if os.path.isfile(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
                data["_meta_path"] = path
                return data
    except Exception:
        pass
    return {"_meta_path": path}

# ---- Primary order lookup (rôzne schémy) -----------------------------------
def _sms_get_b2c_order_info(order_id: Optional[int] = None, order_no: Optional[str] = None) -> dict:
    candidates = ("b2c_objednavky", "b2c_orders", "objednavky_b2c", "objednavky")
    for t in candidates:
        if not _sms_table_exists(t): continue
        cols = _sms_cols_for(t)
        c_id   = _sms_pick(cols, "id", "order_id")
        c_no   = _sms_pick(cols, "cislo_objednavky", "order_number", "objednavka_cislo", "cislo")
        c_mail = _sms_pick(cols, "email", "mail", "zakaznik_email", "customer_email", "email_zakaznika")
        c_tel  = _sms_pick(cols, "telefon", "phone", "tel", "mobil", "mobilne_cislo")
        c_sum  = _sms_pick(cols, "celkova_suma_s_dph", "celkova_suma", "finalna_suma", "final_price",
                                  "suma_s_dph", "sum_celkom", "celkom", "amount_due")
        c_cust = _sms_pick(cols, "zakaznik_id", "customer_id", "b2c_customer_id", "user_id", "customer_uuid", "customer_ref")

        where, params = [], []
        if order_id and c_id: where.append(f"o.{c_id}=%s"); params.append(int(order_id))
        if order_no and c_no: where.append(f"o.{c_no}=%s"); params.append(order_no)
        if not where: continue

        row = db_connector.execute_query(
            f"SELECT {c_no or 'NULL'} AS order_no, {c_mail or 'NULL'} AS email, "
            f"{c_tel or 'NULL'} AS phone, {c_sum or 'NULL'} AS final_price, {c_cust or 'NULL'} AS customer_ref "
            f"FROM {t} o WHERE {' OR '.join(where)} LIMIT 1",
            tuple(params), fetch="one"
        )
        if row:
            return {
                "order_no": (row.get("order_no") or None),
                "email": (row.get("email") or None),
                "phone": (row.get("phone") or None),
                "final_price": row.get("final_price"),
                "customer_ref": (None if row.get("customer_ref") in (None,"") else str(row["customer_ref"]))
            }
    return {"order_no": order_no, "email": None, "phone": None, "final_price": None, "customer_ref": None}

# ---- Lookup v zákazníkoch podľa e-mailu / ID -------------------------------
def _sms_lookup_b2c_phone_by_email(email: Optional[str]) -> Optional[str]:
    if not email: return None
    for t in ("b2b_zakaznici","b2c_zakaznici","b2c_customers","customers","zakaznici","users"):
        if not _sms_table_exists(t): continue
        cols = _sms_cols_for(t)
        c_mail = _sms_pick(cols, "email","mail","email_zakaznika","kontakt_email")
        c_tel  = _sms_pick(cols, "telefon","phone","mobil","mobilne_cislo","tel","phone_number","kontakt_telefon")
        if not (c_mail and c_tel): continue
        row = db_connector.execute_query(
            f"SELECT {c_tel} AS phone FROM {t} WHERE LOWER({c_mail})=LOWER(%s) LIMIT 1",
            (email,), fetch="one"
        )
        if row and row.get("phone"):
            return (row["phone"] or "").strip()
    return None

def _sms_lookup_b2c_phone_by_customer_ref(ref: Optional[str]) -> Optional[str]:
    if not ref: return None
    for t in ("b2b_zakaznici","b2c_zakaznici","b2c_customers","customers","zakaznici","users"):
        if not _sms_table_exists(t): continue
        cols = _sms_cols_for(t)
        idcol = _sms_pick(cols, "zakaznik_id","customer_id","user_id","uuid","id","customer_uuid","customer_ref")
        tel   = _sms_pick(cols, "telefon","phone","mobil","mobilne_cislo","tel")
        if not (idcol and tel): continue
        row = db_connector.execute_query(
            f"SELECT {tel} AS phone FROM {t} WHERE {idcol}=%s LIMIT 1",
            (ref,), fetch="one"
        )
        if row and row.get("phone"):
            return (row["phone"] or "").strip()
    return None

# ---- Robust fallback JOIN na b2c_objednavky × b2b_zakaznici -----------------
def _sms_fallback_order_join(order_no: Optional[str]) -> dict:
    """
    Robustný fallback:
      - autodetekuje stĺpce v b2c_objednavky
      - skúsi exact match aj LIKE, aj 14-ciferný suffix
      - JOIN na b2b_zakaznici pre telefón
    """
    import re
    if not order_no or not _sms_table_exists("b2c_objednavky"):
        return {}

    cols_o = _sms_cols_for("b2c_objednavky")
    c_no    = _sms_pick(cols_o, "cislo_objednavky","order_number","objednavka_cislo","cislo")
    c_cust  = _sms_pick(cols_o, "zakaznik_id","customer_id","user_id")
    c_email = _sms_pick(cols_o, "email","mail","email_zakaznika")
    c_sum   = _sms_pick(cols_o, "celkova_suma_s_dph","celkova_suma","finalna_suma","final_price",
                                  "suma_s_dph","sum_celkom","celkom","amount_due")
    if not c_no:
        return {}

    join_sql = ""
    sel = [f"o.{c_no} AS order_no"]
    sel.append(f"o.{c_email} AS email" if c_email else "NULL AS email")
    sel.append(f"o.{c_cust} AS customer_ref" if c_cust else "NULL AS customer_ref")
    sel.append(f"o.{c_sum} AS final_price" if c_sum else "NULL AS final_price")

    if _sms_table_exists("b2b_zakaznici") and c_cust:
        cols_z = _sms_cols_for("b2b_zakaznici")
        z_id   = _sms_pick(cols_z, "zakaznik_id","customer_id","user_id","id")
        z_tel  = _sms_pick(cols_z, "telefon","phone","mobil","mobilne_cislo","tel")
        if z_id and z_tel:
            sel.append(f"z.{z_tel} AS phone")
            join_sql = f" LEFT JOIN b2b_zakaznici z ON z.{z_id} = o.{c_cust} "
        else:
            sel.append("NULL AS phone")
    else:
        sel.append("NULL AS phone")

    base_sql = f"SELECT {', '.join(sel)} FROM b2c_objednavky o {join_sql} WHERE {{COND}} LIMIT 1"

    # exact
    row = db_connector.execute_query(base_sql.replace("{COND}", f"o.{c_no}=%s"), (order_no,), fetch="one") or {}
    if row and (row.get("order_no") or row.get("phone") or row.get("customer_ref")):
        for k in ("order_no","email","customer_ref","phone"):
            if row.get(k) is not None: row[k] = (str(row[k]) or "").strip()
        return row

    # like (full)
    row = db_connector.execute_query(base_sql.replace("{COND}", f"o.{c_no} LIKE %s"), (f"%{order_no}%",), fetch="one") or {}
    if row and (row.get("order_no") or row.get("phone") or row.get("customer_ref")):
        for k in ("order_no","email","customer_ref","phone"):
            if row.get(k) is not None: row[k] = (str(row[k]) or "").strip()
        return row

    # like (14-digit suffix)
    m = re.search(r"(\d{14})", str(order_no))
    if m:
        suf = m.group(1)
        row = db_connector.execute_query(base_sql.replace("{COND}", f"o.{c_no} LIKE %s"), (f"%{suf}%",), fetch="one") or {}
        if row and (row.get("order_no") or row.get("phone") or row.get("customer_ref")):
            for k in ("order_no","email","customer_ref","phone"):
                if row.get(k) is not None: row[k] = (str(row[k]) or "").strip()
            return row

    return {}

# ---- Kompozit resolve -------------------------------------------------------
def _sms_resolve_b2c_phone(order_id, order_no, phone_hint, email_hint, cust_ref):
    ms = _sms_normalize_msisdn(phone_hint)
    if ms: return ms

    info = _sms_get_b2c_order_info(order_id, order_no)
    ms = _sms_normalize_msisdn(info.get("phone"))
    if ms: return ms

    meta = _sms_read_order_meta(order_no or info.get("order_no") or "")
    ms = _sms_normalize_msisdn(meta.get("phone"))
    if ms: return ms

    ref = cust_ref or info.get("customer_ref")
    ms = _sms_normalize_msisdn(_sms_lookup_b2c_phone_by_customer_ref(ref))
    if ms: return ms

    email = (email_hint or info.get("email"))
    ms = _sms_normalize_msisdn(_sms_lookup_b2c_phone_by_email(email))
    return ms

# ---- _sms_send musí byť NAD notify funkciami (Pylance) ----------------------
def _sms_send(text: str, msisdn: Optional[str]) -> dict:
    if not _sms_mod:
        return {"id": "NO_SMS_MODULE", "note": "sms_handler nie je načítaný"}
    ms = _sms_normalize_msisdn(msisdn)
    if not ms:
        return {"id": "NO_MSISDN", "note": "Chýba alebo neplatné telefónne číslo"}
    return _sms_mod.send_batch(message=text, recipients=[ms], simple_text=True)

# ---- Anti-dup lock (bez DB) -------------------------------------------------
_SMS_LOCK_DIR = os.path.join(os.path.dirname(__file__), "static", "uploads", "sms", "locks")
os.makedirs(_SMS_LOCK_DIR, exist_ok=True)

def _sms_sanitize(s: str) -> str:
    return "".join(ch for ch in str(s or "") if ch.isalnum() or ch in ("-","_"))

def _sms_lock_path(order_no: str, kind: str) -> str:
    return os.path.join(_SMS_LOCK_DIR, f"{_sms_sanitize(order_no)}.{kind}.lock")

def _sms_try_lock(order_no: Optional[str], kind: str, ttl_sec: int = 60) -> Optional[str]:
    if not order_no:
        return None
    path = _sms_lock_path(order_no, kind)
    now = time.time()
    try:
        fd = os.open(path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
        try:
            os.write(fd, str(now).encode("utf-8"))
        finally:
            os.close(fd)
        return path
    except FileExistsError:
        try:
            age = now - os.path.getmtime(path)
            if age > ttl_sec:
                os.remove(path)
                return _sms_try_lock(order_no, kind, ttl_sec)
        except Exception:
            pass
        return None

def _sms_unlock(lock_path: Optional[str]):
    if not lock_path: return
    try:
        os.remove(lock_path)
    except Exception:
        pass

# ---- Serverové NOTIFY (READY/COMPLETED/POINTS) ------------------------------
def b2c_notify_sms_ready(data: dict):
    """
    Pošle READY SMS (jedna jediná – s idempotenciou). Krátky text bez diakritiky.
    Vstup: { "order_id": int, "order_no": str?, "final_price": float? }
    """
    od = data or {}
    info = _sms_get_b2c_order_info(od.get("order_id"), od.get("order_no"))
    order_no = info.get("order_no") or (od.get("order_no") or "")
    if not order_no:
        return {"error": "Chýba číslo objednávky."}

    meta = _sms_read_order_meta(order_no)
    mp = meta.get("_meta_path")
    if meta.get("ready_sms_sent_at"):
        return {"id":"SKIPPED_DUP","note":"READY SMS uz odoslana","order_no":order_no}

    # finalna suma a body
    fp = od.get("final_price")
    if fp is None:
        fp = info.get("final_price")
    try:
        fp = float(str(fp).replace(",", ".")) if fp is not None else None
    except Exception:
        fp = None
    pts = int(fp) if isinstance(fp, (int, float)) and fp > 0 else None

    phone = _sms_resolve_b2c_phone(None, order_no, od.get("phone"), od.get("user_email"), info.get("customer_ref"))
    if not phone:
        return {"error":"Cislo zakaznika sa nenaslo","order_no":order_no}

    # SMS text (<=160)
    base = f"MIK: Objednavka {order_no} je pripravena na vyzdvihnutie."
    if isinstance(fp,(int,float)):
        base += f" Suma {fp:.2f} EUR."
    if isinstance(pts,int) and pts>0:
        base += f" Dakujeme! Po uzavreti pripiseme {pts} bodov."
    lock = _sms_try_lock(order_no, "ready", ttl_sec=30)
    if not lock:
        return {"id":"SKIPPED_LOCK","note":"READY dup <30s","order_no":order_no}
    try:
        res = _sms_send(base, phone)
        try:
            meta["ready_sms_sent_at"] = datetime.utcnow().isoformat() + "Z"
            with open(mp,"w",encoding="utf-8") as f:
                json.dump(meta, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        return res
    finally:
        _sms_unlock(lock)


def b2c_notify_sms_completed(data: dict):
    """
    ZÁMERNE neposielame SMS „objednavka uzavreta…“ – požiadavka: neposielať.
    Vrátime len úspech, aby UI nehlásilo chybu.
    """
    od = data or {}
    return {"id": "SKIPPED", "note": "Completed SMS vypnutá", "order_no": (od.get("order_no") or None)}


def b2c_notify_sms_points(data: dict):
    """
    Pošle SMS o pripísaných bodoch v tvare:
      'Dakujeme za vas nakup! Na Vas ucet sme pripisali XYZ vernostnych bodov!'
    Idempotentné cez META flag 'points_sms_sent_at'.
    """
    od = data or {}
    order_no = od.get("order_no") or ""
    meta = _sms_read_order_meta(order_no) if order_no else {"_meta_path": None}

    if meta.get("points_sms_sent_at"):
        return {"id":"ALREADY_SENT","note":"POINTS SMS uz odoslana","order_no":order_no or None}

    lock = _sms_try_lock(order_no or (od.get("user_email") or od.get("phone") or "points"), "points", ttl_sec=30)
    if not lock:
        return {"id":"SKIPPED_DUP","note":"POINTS dup <30s","order_no":order_no or None}
    try:
        pts = int(od.get("points_delta") or od.get("points_added") or 0)
        if pts <= 0:
            return {"id":"ZERO_POINTS","note":"Bez bodov","order_no":order_no or None}

        phone = _sms_resolve_b2c_phone(od.get("order_id"), order_no, od.get("phone"), od.get("user_email"), od.get("customer_ref"))
        text = f"Dakujeme za vas nakup! Na Vas ucet sme pripisali {pts} vernostnych bodov!"
        out = _sms_send(text, phone)

        try:
            mp = meta.get("_meta_path")
            if mp:
                meta["points_sms_sent_at"] = datetime.utcnow().isoformat() + "Z"
                with open(mp, "w", encoding="utf-8") as f:
                    json.dump(meta, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        return out
    finally:
        _sms_unlock(lock)

# ---- ERP SMS Connector ------------------------------------------------------
def sms_get_recipients():
    from flask import request
    import re
    def _normalize_local(n: str, cc: str = "421") -> Optional[str]:
        s = re.sub(r"[^\d+]", "", str(n or ""))
        if not s: return None
        if s.startswith("00"): s = "+" + s[2:]
        if s.startswith("+"):  pass
        elif s.startswith("0") and len(s) >= 9: s = f"+{cc}{s[1:]}"
        elif len(s) == 9 and s[0] in "9": s = f"+{cc}{s}"
        else: return None
        return s

    limit = int((request.args.get("limit") or "500")); limit = max(1, min(limit, 5000))
    marketing_only = str(request.args.get("marketing_only","1")).lower() not in ("0","false","no")
    q = (request.args.get("q") or "").strip()

    for t in ("b2b_zakaznici","b2c_zakaznici","b2c_customers","customers","zakaznici","users"):
        if _sms_table_exists(t):
            chosen = t; break
    else:
        return []

    colset = _sms_cols_for(chosen)
    def pick(*names, default=None):
        for n in names:
            if n in colset: return n
        return default

    col_id   = pick("zakaznik_id","id","user_id","customer_id","uuid")
    col_name = pick("nazov_firmy","full_name","meno","name","first_name","nick")
    col_mail = pick("email","mail","email_zakaznika")
    col_tel  = pick("telefon","phone","mobil","mobilne_cislo","tel","phone_number")
    col_mkt  = pick("marketing_sms","marketing_suhlas_sms","gdpr_marketing_sms","marketing","newsletter","marketing_consent","newsletter_consent")
    if not col_tel: return []

    where = [f"TRIM(COALESCE({col_tel},''))<>''"]; params = []
    if marketing_only and col_mkt: where.append(f"COALESCE({col_mkt},0)<>0")
    if q:
        ql = f"%{q.lower()}%"; like = []
        if col_name: like.append(f"LOWER(COALESCE({col_name},'')) LIKE %s")
        if col_mail: like.append(f"LOWER(COALESCE({col_mail},'')) LIKE %s")
        like.append(f"LOWER(COALESCE({col_tel},'')) LIKE %s")
        where.append("(" + " OR ".join(like) + ")")
        if col_name: params.append(ql)
        if col_mail: params.append(ql)
        params.append(ql)

    sql = f"SELECT {col_id or 'NULL'} AS id, {col_name or 'NULL'} AS name, {col_tel} AS phone, {col_mkt or 'NULL'} AS marketing_sms FROM {chosen} WHERE {' AND '.join(where)} ORDER BY {col_name or col_tel} ASC LIMIT {limit}"
    rows = db_connector.execute_query(sql, tuple(params)) or []
    out = []
    for r in rows:
        raw = (r.get("phone") or "").strip()
        msisdn = _normalize_local(raw)
        if not msisdn: continue
        out.append({"id": r.get("id"), "name": r.get("name") or "", "phone": raw, "msisdn": msisdn, "marketing_sms": bool(r.get("marketing_sms")) if col_mkt else None})
    return out

def sms_send(data: dict):
    import re, os
    if not _sms_mod:
        return {"error":"SMS modul nie je načítaný (chýba sms_handler.py)."}
    msg = (data or {}).get("message") or ""
    if not msg.strip():
        return {"error":"Zadajte text správy."}
    rec_in = (data or {}).get("recipients") or []
    recs = []
    if isinstance(rec_in, list):
        for it in rec_in:
            if isinstance(it, str):
                ms = _sms_mod.normalize_msisdn(it)
            elif isinstance(it, dict):
                ms = _sms_mod.normalize_msisdn(it.get("msisdn") or it.get("phone"))
            else:
                ms = None
            if ms: recs.append(ms)
    recs = sorted(set(recs))
    if not recs:
        return {"error":"Žiadni príjemcovia (nepodarilo sa normalizovať čísla)."}
    sender = (data or {}).get("sender") or os.getenv("SMS_SENDER_ID") or "MIK"
    sender = re.sub(r"[^A-Z0-9._-]","",str(sender or "").upper())[:11] or "MIK"
    simple_text = bool((data or {}).get("simple_text", True))
    res = _sms_mod.send_batch(message=msg, recipients=recs, sender=sender, simple_text=simple_text,
                              department=(data or {}).get("department"), schedule=(data or {}).get("schedule"),
                              callback_url=(data or {}).get("callback_url"))
    if res.get("id") != "OK":
        return {"error":res.get("note") or res.get("id") or "Odoslanie zlyhalo", "response":res}
    acc = (res.get("data") or {}).get("recipients", {}).get("accepted", []) or []
    return {"message":f"Odoslané {len(acc)} / {len(recs)}", "batch_id":(res.get("data") or {}).get("batch_id"), "accepted":acc}

# ---- Diagnostika / náhľad ---------------------------------------------------
_SMS_DEBUG_DIR = os.path.join(os.path.dirname(__file__), "static", "uploads", "sms")
os.makedirs(_SMS_DEBUG_DIR, exist_ok=True)
_SMS_DEBUG_LOG = os.path.join(_SMS_DEBUG_DIR, "sms_debug.jsonl")

def _dbg_write(event: str, payload: dict):
    try:
        rec = {"event":event, **payload}
        with open(_SMS_DEBUG_LOG, "a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception:
        pass

def _cand_tables_exist() -> dict:
    out = {}
    for t in ("b2c_objednavky","b2c_orders","objednavky_b2c","objednavky","b2b_zakaznici","b2c_zakaznici","b2c_customers","customers","zakaznici","users"):
        if _sms_table_exists(t):
            out[t] = _sms_cols_for(t)
    return out

def _lookup_from_order(tables: dict, order_id: Optional[int], order_no: Optional[str]) -> dict:
    checked = []
    for t in ("b2c_objednavky","b2c_orders","objednavky_b2c","objednavky"):
        if t not in tables: continue
        cols = tables[t]
        c_id   = _sms_pick(cols, "id","order_id")
        c_no   = _sms_pick(cols, "cislo_objednavky","order_number","objednavka_cislo","cislo")
        c_mail = _sms_pick(cols, "email","mail","zakaznik_email","customer_email","email_zakaznika")
        c_tel  = _sms_pick(cols, "telefon","phone","tel","mobil","mobilne_cislo")
        c_sum  = _sms_pick(cols, "celkova_suma_s_dph","celkova_suma","finalna_suma","final_price",
                                  "suma_s_dph","sum_celkom","celkom","amount_due")
        c_cust = _sms_pick(cols, "zakaznik_id","customer_id","b2c_customer_id","user_id","customer_uuid","customer_ref")
        where, params = [], []
        if order_id and c_id: where.append(f"{c_id}=%s"); params.append(int(order_id))
        if order_no and c_no: where.append(f"{c_no}=%s"); params.append(order_no)
        if not where: continue
        row = db_connector.execute_query(
            f"SELECT {c_no or 'NULL'} AS order_no, {c_mail or 'NULL'} AS email, {c_tel or 'NULL'} AS phone, "
            f"{c_sum or 'NULL'} AS final_price, {c_cust or 'NULL'} AS customer_ref "
            f"FROM {t} WHERE {' OR '.join(where)} LIMIT 1",
            tuple(params), fetch="one"
        )
        checked.append({"table": t, "where": where, "row": row})
        if row:
            return {"found_in": t, "row": row, "checked": checked}
    return {"found_in": None, "row": None, "checked": checked}

def _lookup_from_customer_tables(tables: dict, email: Optional[str], customer_ref: Optional[str]) -> dict:
    """Skúsi nájsť telefón v zákazníckych tabuľkách – najprv podľa customer_ref, potom podľa e-mailu."""
    checked = []
    if customer_ref:
        for t in ("b2b_zakaznici","b2c_zakaznici","b2c_customers","customers","zakaznici","users"):
            if t not in tables: continue
            cols = tables[t]
            idcol = _sms_pick(cols, "zakaznik_id","customer_id","user_id","uuid","id","customer_uuid","customer_ref")
            tel   = _sms_pick(cols, "telefon","phone","mobil","mobilne_cislo","tel")
            if not (idcol and tel): continue
            row = db_connector.execute_query(
                f"SELECT {tel} AS phone FROM {t} WHERE {idcol}=%s LIMIT 1",
                (customer_ref,), fetch="one"
            )
            checked.append({"table": t, "by": "customer_ref", "row": row})
            if row and row.get("phone"):
                return {"found_in": t, "row": row, "checked": checked}
    if email:
        for t in ("b2b_zakaznici","b2c_zakaznici","b2c_customers","customers","zakaznici","users"):
            if t not in tables: continue
            cols = tables[t]
            c_mail = _sms_pick(cols, "email","mail","email_zakaznika","kontakt_email")
            tel    = _sms_pick(cols, "telefon","phone","mobil","mobilne_cislo","tel","phone_number","kontakt_telefon")
            if not (c_mail and tel): continue
            row = db_connector.execute_query(
                f"SELECT {tel} AS phone FROM {t} WHERE LOWER({c_mail})=LOWER(%s) LIMIT 1",
                (email,), fetch="one"
            )
            checked.append({"table": t, "by": "email", "row": row})
            if row and row.get("phone"):
                return {"found_in": t, "row": row, "checked": checked}
    return {"found_in": None, "row": None, "checked": checked}

def b2c_sms_where_phone(data: dict):
    od = data or {}
    order_id     = od.get("order_id")
    order_no     = (od.get("order_no") or "").strip()
    user_email   = od.get("user_email") or od.get("email")
    customer_ref = od.get("customer_ref")
    phone_hint   = od.get("phone")

    tables = _cand_tables_exist()
    hint_msisdn = _sms_normalize_msisdn(phone_hint)

    oinfo = _lookup_from_order(tables, order_id, order_no)
    order_row = (oinfo.get("row") or {}) if oinfo else {}

    if not order_row or not order_row.get("order_no"):
        fb = _sms_fallback_order_join(order_no)
        if fb and (fb.get("order_no") or fb.get("phone") or fb.get("customer_ref")):
            order_row = {"order_no":fb.get("order_no"), "email":fb.get("email"), "phone":fb.get("phone"),
                         "final_price":fb.get("final_price"), "customer_ref":fb.get("customer_ref")}

    meta = _sms_read_order_meta(order_no or order_row.get("order_no"))
    from_meta = meta.get("phone")

    cinfo = _lookup_from_customer_tables(tables, user_email or order_row.get("email"), customer_ref or order_row.get("customer_ref"))
    cust_row = (cinfo.get("row") or {}) if cinfo else {}

    from_order    = order_row.get("phone")
    from_customer = cust_row.get("phone")
    resolved = (hint_msisdn
                or _sms_normalize_msisdn(from_order)
                or _sms_normalize_msisdn(from_meta)
                or _sms_normalize_msisdn(from_customer))

    out = {
        "search_params":{"order_id":order_id,"order_no":order_no or order_row.get("order_no"),
                         "user_email":user_email,"customer_ref":customer_ref,"phone_hint":phone_hint},
        "order_lookup":{"source":(oinfo or {}).get("found_in"),"row":order_row,"checked":(oinfo or {}).get("checked")},
        "meta_lookup":{"path":meta.get("_meta_path"),"phone":from_meta},
        "customer_lookup":{"source":(cinfo or {}).get("found_in"),"row":cust_row,"checked":(cinfo or {}).get("checked")},
        "resolved_msisdn":resolved
    }
    _dbg_write("wherePhone", out)
    return out

def b2c_sms_diag(data: dict):
    od = data or {}
    where = b2c_sms_where_phone(od)
    msisdn = where.get("resolved_msisdn")
    typ = (od.get("type") or "").lower()
    if typ == "ready":
        fp = od.get("final_price")
        if fp is None: fp = (where.get("order_lookup") or {}).get("row",{}).get("final_price")
        txt = f"MIK: objednavka {od.get('order_no') or (where.get('order_lookup') or {}).get('row',{}).get('order_no','')} je pripravena na vyzdvihnutie."
        if fp is not None:
            try: txt += f" Suma {float(fp):.2f} €."
            except Exception: pass
    elif typ == "completed":
        txt = f"MIK: objednavka {od.get('order_no') or (where.get('order_lookup') or {}).get('row',{}).get('order_no','')} uzavreta."
        if od.get("final_paid") is not None:
            try: txt += f" Uhradene {float(od['final_paid']):.2f} €."
            except Exception: pass
        pts = od.get("points_added")
        if isinstance(pts,(int,float)) and pts != 0:
            txt += f" Body {('+' if pts>=0 else '')}{int(pts)}."
    else:
        pts = int(od.get("points_delta") or 0)
        txt = f"MIK: vernostne body zmena {('+' if pts>=0 else '')}{pts}."
    out = {"resolved_msisdn":msisdn,"text":txt,"has_sms_handler":bool(_sms_mod),"wherePhone":where}
    _dbg_write("diag", out)
    return out

def b2c_sms_send_test(data: dict):
    if not _sms_mod:
        return {"error":"sms_handler nie je načítaný."}
    diag = b2c_sms_diag(data)
    ms = diag.get("resolved_msisdn") or _sms_normalize_msisdn((data or {}).get("override_phone"))
    if not ms:
        return {"error":"Nepodarilo sa určiť MSISDN."}, 400
    txt = diag.get("text") or "TEST SMS"
    res = _sms_mod.send_batch(message=txt, recipients=[ms], simple_text=True)
    _dbg_write("sendTest", {"msisdn":ms,"text":txt,"response":res})
    return {"message":"Odoslané" if res.get("id")=="OK" else "NEODOSLANE","response":res,"msisdn":ms,"text":txt}

def b2c_notify_email_ready(data: dict):
    """
    Pošle iba E-MAIL 'Pripravená' (bez SMS). Má idempotenciu cez META flag.
    Vstup: { "order_id": int, "order_no": str?, "final_price": float? }
    """
    od = data or {}
    info = _sms_get_b2c_order_info(od.get("order_id"), od.get("order_no"))
    order_no = info.get("order_no") or (od.get("order_no") or "")
    if not order_no:
        return {"error": "Chýba číslo objednávky."}

    # finálna suma – z requestu alebo z DB/meta
    final_price = od.get("final_price")
    if final_price is None:
        final_price = info.get("final_price")
    try:
        final_price = float(str(final_price).replace(",", ".")) if final_price is not None else None
    except Exception:
        final_price = None

    # e-mail adresa
    to_email = info.get("email")
    if not to_email and info.get("customer_ref"):
        cust = db_connector.execute_query(
            "SELECT email FROM b2b_zakaznici WHERE zakaznik_id=%s OR id=%s LIMIT 1",
            (info["customer_ref"], info["customer_ref"]), fetch="one"
        ) or {}
        to_email = cust.get("email")
    if not to_email:
        return {"error": "E-mail zákazníka sa nenašiel."}

    # idempotencia
    meta = _sms_read_order_meta(order_no)
    mp = meta.get("_meta_path")
    if meta.get("ready_email_sent_at"):
        return {"message": "READY e-mail už bol odoslaný.", "order_no": order_no}

    # odoslanie (text s poďakovaním + info o vyzdvihnutí + predpokladané body)
    try:
        notification_handler.send_b2c_order_ready_email(
            to_email,
            order_no,
            float(final_price) if isinstance(final_price, (int, float)) else None
        )
    except Exception as e:
        return {"error": f"Chyba pri odoslaní e-mailu: {e}"}

    # zapíš flag
    try:
        meta["ready_email_sent_at"] = datetime.utcnow().isoformat() + "Z"
        with open(mp, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

    return {"message": "READY e-mail odoslaný.", "order_no": order_no}



def b2c_notify_email_completed(data: dict):
    """
    Pošle e-mail 'uzavretá' + idempotentne odošle aj SMS o pripísaných bodoch.
    Vstup: { order_id?, order_no?, final_paid?, points_added?, to_email? }
    """
    od = data or {}
    info = _sms_get_b2c_order_info(od.get("order_id"), od.get("order_no"))
    order_no = info.get("order_no") or od.get("order_no") or ""
    if not order_no:
        return {"error": "Chýba číslo objednávky."}

    # e-mail
    to_email = od.get("to_email") or info.get("email")
    if not to_email and info.get("customer_ref"):
        cust = db_connector.execute_query(
            "SELECT email FROM b2b_zakaznici WHERE zakaznik_id=%s OR id=%s LIMIT 1",
            (info["customer_ref"], info["customer_ref"]), fetch="one"
        ) or {}
        to_email = cust.get("email")
    if not to_email:
        return {"error": "E-mail zákazníka sa nenašiel."}

    # sumy/body
    final_paid = od.get("final_paid")
    if final_paid is None:
        final_paid = info.get("final_price")
    try:
        final_paid = float(str(final_paid).replace(",", ".")) if final_paid is not None else 0.0
    except Exception:
        final_paid = 0.0
    pts = od.get("points_added")
    try:
        pts = int(pts) if pts is not None else int(final_paid)
    except Exception:
        pts = 0

    meta = _sms_read_order_meta(order_no)
    if not meta.get("completed_email_sent_at"):
        notification_handler.send_b2c_order_completed_email(to_email, order_no, final_paid, pts)
        try:
            mp = meta.get("_meta_path"); meta["completed_email_sent_at"] = datetime.utcnow().isoformat() + "Z"
            with open(mp, "w", encoding="utf-8") as f:
                json.dump(meta, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    # +++ NOVÉ: po e-maile pošli aj body-SMS (idempotentne)
    try:
        # ak už odišla v minulosti, skip
        if not meta.get("points_sms_sent_at"):
            b2c_notify_sms_points({
                "order_id": od.get("order_id"),
                "order_no": order_no,
                "points_added": pts,
                "customer_ref": info.get("customer_ref"),
                "user_email": to_email
            })
    except Exception as _e:
        print("points SMS after completed email error:", _e)

    return {"message": "COMPLETED e-mail (a body-SMS) spracované.", "order_no": order_no}
def calculate_production_plan():
    """
    Týždenný plán výroby – iba vlastné výrobky z centrálneho katalógu:
    typy VÝROBOK, VÝROBOK_KRAJANY, VÝROBOK_KUSOVY.

    Tabuľky:
      - produkty (nazov_vyrobku, minimalna_zasoba_kg, aktualny_sklad_finalny_kg,
                  kategoria_pre_recepty, predajna_kategoria, vyrobna_davka_kg, ean, vaha_balenia_g, + typ*)
      - b2b_objednavky + b2b_objednavky_polozky

    Výstup (pre planning.js): { "Kategória": [ {nazov_vyrobku, celkova_potreba, aktualny_sklad, navrhovana_vyroba}, ... ], ... }
    """
    import math
    import db_connector

    # Jedna kolácia na porovnania (fix 'Illegal mix of collations')
    COLL = "utf8mb4_0900_ai_ci"
    ALLOWED_TYPES = ("VÝROBOK", "VÝROBOK_KRAJANY", "VÝROBOK_KUSOVY")

    # --- zisti názov stĺpca pre typ produktu v 'produkty' ---
    cols = db_connector.execute_query(
        "SELECT COLUMN_NAME FROM information_schema.COLUMNS "
        "WHERE TABLE_SCHEMA = DATABASE() AND TABLE_NAME = 'produkty'",
        fetch="all"
    ) or []
    colset = {c["COLUMN_NAME"] for c in cols}
    type_col = None
    for cand in ("typ_produktu", "typ", "product_type", "centralny_typ", "typ_katalogu", "kategoria_centralna", "predajna_kategoria"):
        if cand in colset:
            type_col = cand
            break

    # --- načítaj produkty (len povolené typy) ---
    where_type = ""
    if type_col:
        in_list = ", ".join([f"CONVERT('{v}' USING utf8mb4) COLLATE {COLL}" for v in ALLOWED_TYPES])
        where_type = f"WHERE CONVERT(p.{type_col} USING utf8mb4) COLLATE {COLL} IN ({in_list})"

    products_sql = f"""
        SELECT
            p.nazov_vyrobku                                        AS name,
            COALESCE(p.minimalna_zasoba_kg, 0)                     AS min_stock,
            COALESCE(p.aktualny_sklad_finalny_kg, 0)               AS stock,
            COALESCE(p.kategoria_pre_recepty, p.predajna_kategoria, 'Nezaradené') AS cat,
            COALESCE(p.vyrobna_davka_kg, 0)                        AS batch_kg,
            p.ean                                                  AS ean,
            COALESCE(p.vaha_balenia_g, 0)                          AS pack_g
        FROM produkty p
        {where_type}
    """
    products = db_connector.execute_query(products_sql, fetch="all") or []

    # mapa pre lookup a limitovanie len na "tvoje" výrobky
    by_name = {(r["name"] or "").strip(): r for r in products}

    # --- dopyt z otvorených objednávok (kolácie na porovnaniach) ---
    demand_sql = f"""
        SELECT
            COALESCE(
                CONVERT(p.nazov_vyrobku USING utf8mb4) COLLATE {COLL},
                CONVERT(pol.nazov_vyrobku USING utf8mb4) COLLATE {COLL}
            ) AS name,
            SUM(
              CASE
                WHEN LOWER(COALESCE(pol.mj,'')) = 'kg' THEN COALESCE(pol.mnozstvo,0)
                WHEN LOWER(COALESCE(pol.mj,'')) = 'ks' THEN COALESCE(pol.mnozstvo,0)
                      * COALESCE(p.vaha_balenia_g, pol.vaha_balenia_g, 0)/1000
                ELSE COALESCE(pol.mnozstvo,0)
              END
            ) AS q_kg
        FROM b2b_objednavky_polozky pol
        JOIN b2b_objednavky o ON o.id = pol.objednavka_id
        LEFT JOIN produkty p ON (
             (p.ean IS NOT NULL AND pol.ean_produktu IS NOT NULL
               AND CONVERT(p.ean USING utf8mb4) COLLATE {COLL}
                   = CONVERT(pol.ean_produktu USING utf8mb4) COLLATE {COLL})
          OR (CONVERT(p.nazov_vyrobku USING utf8mb4) COLLATE {COLL}
                   = CONVERT(pol.nazov_vyrobku USING utf8mb4) COLLATE {COLL})
        )
        WHERE COALESCE(CONVERT(o.stav USING utf8mb4) COLLATE {COLL},'') NOT IN
              ('Zrušená','Zrusena','Zrušena','Hotová','Hotova','Uhradená','Uhradena')
        GROUP BY name
    """
    demand_rows = db_connector.execute_query(demand_sql, fetch="all") or []

    # dopyt mapujeme, ale v pláne použijeme IBA produkty, ktoré sú v 'by_name' (tvoje typy)
    demand_map = {}
    for r in demand_rows:
        n = (r.get("name") or "").strip()
        if not n or n not in by_name:
            continue
        try:
            q = float(r.get("q_kg") or 0.0)
        except Exception:
            q = 0.0
        demand_map[n] = demand_map.get(n, 0.0) + max(q, 0.0)

    # --- zostav výstup po kategóriách, len pre tvoje výrobky ---
    out = {}
    for name, prod in by_name.items():
        cat       = prod.get("cat") or "Nezaradené"
        min_stock = float(prod.get("min_stock") or 0)
        stock     = float(prod.get("stock") or 0)
        batch_kg  = float(prod.get("batch_kg") or 0)
        demand    = float(demand_map.get(name, 0.0))

        need_for_min = max(min_stock - stock, 0.0)
        total_need   = max(demand, 0.0) + need_for_min

        raw_make = max(max(min_stock, demand) - stock, 0.0)
        make = math.ceil(raw_make / batch_kg) * batch_kg if (batch_kg > 0 and raw_make > 0) else raw_make

        item = {
            "nazov_vyrobku": name,
            "celkova_potreba": round(total_need, 3),
            "aktualny_sklad": round(stock, 3),
            "navrhovana_vyroba": round(make, 3),
        }
        out.setdefault(cat, []).append(item)

    # zoradenie v rámci kategórií
    for cat, items in out.items():
        items.sort(key=lambda x: (-x["navrhovana_vyroba"], -x["celkova_potreba"], x["nazov_vyrobku"]))

    return out

    # ── helpers ─────────────────────────────────────────────────
    def _tbl_exists(t):
        row = db_connector.execute_query(
            "SELECT COUNT(*) AS c FROM information_schema.TABLES WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME=%s",
            (t,), fetch="one")
        return bool(row and int(list(row.values())[0]) > 0)

    def _cols(t):
        rows = db_connector.execute_query(
            "SELECT COLUMN_NAME FROM information_schema.COLUMNS WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME=%s",
            (t,), fetch="all") or []
        return {r["COLUMN_NAME"] for r in rows}

    def _pick(colset, *cands):
        for c in cands:
            if c in colset:
                return c
        return None

    # ── 1) Meta výrobkov: min. zásoba + kategória ──────────────
    product_meta = {}  # name -> (min_stock_kg, category)
    prod_tables = ["vyrobky", "produkty", "hotove_vyrobky", "b2c_vyrobky"]
    for t in prod_tables:
        if not _tbl_exists(t):
            continue
        cs = _cols(t)
        name_col = _pick(cs, "nazov_vyrobku", "nazov", "produkt", "produkt_nazov", "name")
        min_col  = _pick(cs, "min_zasoba_kg", "min_zasoba", "min_sklad_kg", "min_sklad")
        cat_col  = _pick(cs, "kategoria", "category", "typ", "skupina")
        if not name_col:
            continue
        rows = db_connector.execute_query(
            f"SELECT {name_col} AS n, {min_col} AS mn, {cat_col} AS cat FROM {t}",
            fetch="all") or []
        for r in rows:
            n = (r.get("n") or "").strip()
            if not n:
                continue
            mn = 0.0
            try:
                mn = float(r.get("mn") or 0)
            except Exception:
                mn = 0.0
            cat = r.get("cat") or "Nezaradené"
            product_meta[n] = (mn, cat)
        if product_meta:
            break  # stačí prvá nájdená tabuľka

    # ── 2) Sklad hotových výrobkov ─────────────────────────────
    stock_map = {}  # name -> kg
    stock_tables = ["sklad_vyrobkov", "sklad_hotove", "sklad_produkty", "sklad"]
    for t in stock_tables:
        if not _tbl_exists(t):
            continue
        cs = _cols(t)
        name_col = _pick(cs, "nazov_vyrobku", "nazov", "produkt", "produkt_nazov", "name")
        qty_col  = _pick(cs, "mnozstvo_kg", "stav_kg", "mnozstvo", "stav", "qty_kg", "qty")
        if not (name_col and qty_col):
            continue
        rows = db_connector.execute_query(
            f"SELECT {name_col} AS n, SUM({qty_col}) AS q FROM {t} GROUP BY {name_col}", fetch="all") or []
        for r in rows:
            n = (r.get("n") or "").strip()
            if not n:
                continue
            q = 0.0
            try:
                q = float(r.get("q") or 0)
            except Exception:
                q = 0.0
            stock_map[n] = stock_map.get(n, 0.0) + q
        if stock_map:
            break

    # ── 3) Dopyt z neuzavretých objednávok ─────────────────────
    demand = {}  # name -> kg
    order_item_candidates = [
        ("b2c_objednavky_polozky", "b2c_objednavky", "objednavka_id"),
        ("b2b_objednavky_polozky", "b2b_objednavky", "objednavka_id"),
        ("objednavky_polozky", "objednavky", "objednavka_id"),
    ]
    for itbl, otbl, fk_guess in order_item_candidates:
        if not _tbl_exists(itbl):
            continue
        ics = _cols(itbl)
        name_col = _pick(ics, "nazov_vyrobku", "nazov", "produkt_nazov", "polozka_nazov", "vyrobok_nazov", "produkt")
        qty_col  = _pick(ics, "mnozstvo_kg", "mnozstvo", "qty_kg", "qty")
        if not (name_col and qty_col):
            continue

        join = ""
        where = []
        if _tbl_exists(otbl):
            ocs = _cols(otbl)
            ofk = _pick(ics, fk_guess, "objednavka_id", "order_id")
            oid = "id" if "id" in ocs else None
            if ofk and oid:
                join = f" JOIN {otbl} o ON o.{oid} = {itbl}.{ofk}"
                st_col = "stav" if "stav" in ocs else ("status" if "status" in ocs else None)
                if st_col:
                    where.append(f"COALESCE(o.{st_col}, '') NOT IN ('Zrušená','Zrusena','Zrušena','Hotová','Hotova','Uhradená','Uhradena')")
        sql = f"SELECT {name_col} AS n, SUM({qty_col}) AS q FROM {itbl}{join}"
        if where:
            sql += " WHERE " + " AND ".join(where)
        sql += f" GROUP BY {name_col}"

        rows = db_connector.execute_query(sql, fetch="all") or []
        for r in rows:
            n = (r.get("n") or "").strip()
            if not n:
                continue
            q = 0.0
            try:
                q = float(r.get("q") or 0)
            except Exception:
                q = 0.0
            demand[n] = demand.get(n, 0.0) + q

    # ── 4) Zostavenie výsledku ──────────────────────────────────
    items_by_cat = {}
    names = set(product_meta.keys()) | set(demand.keys()) | set(stock_map.keys())
    for name in sorted(names):
        min_stock, cat = product_meta.get(name, (0.0, "Nezaradené"))
        stock  = float(stock_map.get(name, 0.0))
        dem    = max(float(demand.get(name, 0.0)), 0.0)

        target = max(min_stock, dem)
        to_make = max(target - stock, 0.0)
        total_need = dem + max(min_stock - stock, 0.0)

        item = {
            "nazov_vyrobku": name,
            "celkova_potreba": round(total_need, 3),
            "aktualny_sklad":  round(stock, 3),
            "navrhovana_vyroba": round(to_make, 3),
        }
        items_by_cat.setdefault(cat or "Nezaradené", []).append(item)

    return items_by_cat
def get_purchase_suggestions():
    """
    Návrh nákupu surovín:
      - zoberie 'navrhovana_vyroba' z calculate_production_plan(),
      - dopočíta spotrebu surovín podľa receptov (recepty.mnozstvo_na_davku_kg),
        s prepočtom na 1 kg (delené vyrobna_davka_kg z 'produkty'),
      - odpočíta aktuálny stav surovín zo 'sklad' a dorovná na min_zasoba.
    Výstup (planning.js): [ { name, currentStock, requiredForProduction, minStock, purchaseQty }, ... ]
    """
    import db_connector

    # 1) Množstvá plánovanej výroby podľa produktu (kg)
    try:
        plan = calculate_production_plan() or {}
    except Exception:
        plan = {}
    plan_qty = {}  # vyrobok -> kg
    for items in plan.values():
        for it in (items or []):
            if not it:
                continue
            n = (it.get("nazov_vyrobku") or "").strip()
            q = 0.0
            try:
                q = float(it.get("navrhovana_vyroba") or 0)
            except Exception:
                q = 0.0
            if n and q > 0:
                plan_qty[n] = plan_qty.get(n, 0.0) + q

    if not plan_qty:
        # netreba dokupovať – žiadny plán
        return []

    # 2) Načítaj vyrobna_davka_kg pre prepočet receptov na 1 kg
    prod_rows = db_connector.execute_query("""
        SELECT nazov_vyrobku AS n, COALESCE(vyrobna_davka_kg,0) AS batch_kg
        FROM produkty
    """, fetch="all") or []
    batch_map = { (r["n"] or "").strip(): float(r.get("batch_kg") or 0) for r in prod_rows }

    # 3) Recepty – koľko suroviny ide na 1 kg výrobku
    rec_rows = db_connector.execute_query("""
        SELECT nazov_vyrobku AS p, nazov_suroviny AS s, COALESCE(mnozstvo_na_davku_kg,0) AS per_batch
        FROM recepty
    """, fetch="all") or []

    required = {}  # surovina -> kg potrebné
    for r in rec_rows:
        p = (r.get("p") or "").strip()
        s = (r.get("s") or "").strip()
        if not p or not s:
            continue
        q_plan = plan_qty.get(p, 0.0)
        if q_plan <= 0:
            continue
        per_batch = float(r.get("per_batch") or 0.0)
        batch_kg  = float(batch_map.get(p, 0.0))
        # koeficient na 1 kg: ak nemáme vyrobna_davka_kg, berieme per_batch ako per-kg (fallback)
        per_kg = per_batch / batch_kg if batch_kg > 0 else per_batch
        need = q_plan * max(per_kg, 0.0)
        if need > 0:
            required[s] = required.get(s, 0.0) + need

    # 4) Stav a minimá surovín zo 'sklad'
    stock_rows = db_connector.execute_query("""
        SELECT nazov AS n, COALESCE(SUM(mnozstvo),0) AS cur, COALESCE(MAX(min_zasoba),0) AS min_s
        FROM sklad
        GROUP BY nazov
    """, fetch="all") or []
    cur_map  = { (r["n"] or "").strip(): float(r.get("cur") or 0.0)   for r in stock_rows }
    min_map  = { (r["n"] or "").strip(): float(r.get("min_s") or 0.0) for r in stock_rows }

    # 5) Výpočet návrhu nákupu
    out = []
    names = sorted(set(required.keys()) | set(cur_map.keys()) | set(min_map.keys()))
    for name in names:
        req = float(required.get(name, 0.0))
        cur = float(cur_map.get(name, 0.0))
        mn  = float(min_map.get(name, 0.0))
        purchase = max(req + mn - cur, 0.0)
        if purchase <= 0 and req <= 0 and mn <= 0:
            continue
        out.append({
            "name": name,
            "currentStock": round(cur, 3),
            "requiredForProduction": round(req, 3),
            "minStock": round(mn, 3),
            "purchaseQty": round(purchase, 3),
        })

    # zoradenie: najprv tie, čo najviac treba kúpiť
    out.sort(key=lambda x: (-x["purchaseQty"], -x["requiredForProduction"], x["name"]))
    return out


    # ── helpers na introspekciu ─────────────────────────────────
    def _tbl_exists(t):
        row = db_connector.execute_query(
            "SELECT COUNT(*) AS c FROM information_schema.TABLES WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME=%s",
            (t,), fetch="one")
        return bool(row and int(list(row.values())[0]) > 0)

    def _cols(t):
        rows = db_connector.execute_query(
            "SELECT COLUMN_NAME FROM information_schema.COLUMNS WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME=%s",
            (t,), fetch="all") or []
        return {r["COLUMN_NAME"] for r in rows}

    def _pick(colset, *cands):
        for c in cands:
            if c in colset:
                return c
        return None

    # 2) Receptúry – spotreba surovín na 1 kg výrobku
    #    Skúsime viac možných názvov tabuliek/stĺpcov.
    recipe_tables = ["receptury_polozky", "receptura_polozky", "recept_polozky", "recept_ingrediencie"]
    # mapovanie: surovina -> potrebné množstvo (kg) pre plán
    req_map = {}

    for t in recipe_tables:
        if not _tbl_exists(t):
            continue
        cs = _cols(t)
        prod_name = _pick(cs, "vyrobok_nazov", "produkt_nazov", "nazov_vyrobku", "produkt", "vyrobok")
        ingr_name = _pick(cs, "surovina_nazov", "nazov_suroviny", "surovina", "ingredient", "nazov")
        perkg_col = _pick(cs, "mnozstvo_na_kg", "spotreba_na_kg", "na_kg", "koef_na_kg", "mnozstvo")

        if not (prod_name and ingr_name and perkg_col):
            continue

        # načítaj všetky riadky a spočítaj podľa plánu
        rows = db_connector.execute_query(
            f"SELECT {prod_name} AS p, {ingr_name} AS s, {perkg_col} AS k FROM {t}",
            fetch="all") or []
        if not rows:
            continue

        any_used = False
        for r in rows:
            p = (r.get("p") or "").strip()
            s = (r.get("s") or "").strip()
            if not (p and s):
                continue
            coef = 0.0
            try:
                coef = float(r.get("k") or 0)
            except Exception:
                coef = 0.0
            q_plan = plan_qty.get(p, 0.0)
            if q_plan > 0 and coef > 0:
                any_used = True
                req_map[s] = req_map.get(s, 0.0) + q_plan * coef
        if any_used:
            break  # máme použiteľnú tabuľku receptúr

    # 3) Sklad surovín + min. zásoby
    stock_tbls = ["sklad_surovin", "suroviny_sklad", "sklad_surovina"]
    stock = {}  # surovina -> kg
    for t in stock_tbls:
        if not _tbl_exists(t):
            continue
        cs = _cols(t)
        name_col = _pick(cs, "nazov", "surovina", "surovina_nazov", "ingredient")
        qty_col  = _pick(cs, "mnozstvo_kg", "stav_kg", "qty_kg", "mnozstvo", "stav", "qty")
        if not (name_col and qty_col):
            continue
        rows = db_connector.execute_query(
            f"SELECT {name_col} AS n, SUM({qty_col}) AS q FROM {t} GROUP BY {name_col}",
            fetch="all") or []
        for r in rows:
            n = (r.get("n") or "").strip()
            if not n:
                continue
            q = 0.0
            try:
                q = float(r.get("q") or 0)
            except Exception:
                q = 0.0
            stock[n] = stock.get(n, 0.0) + q
        break

    min_tbls = ["suroviny", "materialy", "b2b_suroviny"]
    min_map = {}  # surovina -> min. zásoba (kg)
    for t in min_tbls:
        if not _tbl_exists(t):
            continue
        cs = _cols(t)
        name_col = _pick(cs, "nazov", "surovina", "surovina_nazov", "ingredient", "name")
        min_col  = _pick(cs, "min_zasoba_kg", "min_zasoba", "min_sklad_kg", "min_sklad")
        if not name_col:
            continue
        rows = db_connector.execute_query(
            f"SELECT {name_col} AS n, {min_col} AS m FROM {t}", fetch="all") or []
        for r in rows:
            n = (r.get("n") or "").strip()
            if not n:
                continue
            m = 0.0
            try:
                m = float(r.get("m") or 0)
            except Exception:
                m = 0.0
            min_map[n] = m
        break

    # 4) Zostav výstup
    suggestions = []
    # zahrň suroviny, ktoré treba pre plán (req_map), aj tie, čo sú pod minimom
    names = set(req_map.keys()) | set(min_map.keys()) | set(stock.keys())
    for name in sorted(names):
        required = float(req_map.get(name, 0.0))
        cur      = float(stock.get(name, 0.0))
        min_s    = float(min_map.get(name, 0.0))
        # jednoduchý model: potrebujeme pokryť výrobu + dorovnať na minimálny stav
        purchase = max(required + min_s - cur, 0.0)
        if purchase <= 0 and required <= 0 and min_s <= 0:
            continue
        suggestions.append({
            "name": name,
            "currentStock": round(cur, 3),
            "requiredForProduction": round(required, 3),
            "minStock": round(min_s, 3),
            "purchaseQty": round(purchase, 3),
        })
    return suggestions
def get_7_day_forecast():
    """
    7-dňový prehľad potreby produktov.
    Vracia:
      - dates:            ["YYYY-MM-DD", ... x7]
      - forecast:         B2B + B2C DOKOPY (hlavné pole pre UI)
      - b2c_forecast:     len B2C (na kontrolu / kompatibilitu)
      - forecast_b2c, b2c: aliasy na b2c_forecast
      - debug:            diagnostika (tabuľky/stĺpce/počty)
    """
    from datetime import date, timedelta
    import db_connector

    COLL = "utf8mb4_0900_ai_ci"

    def _dates7():
        base = date.today()
        return [(base + timedelta(days=i)).strftime("%Y-%m-%d") for i in range(7)]

    def _tbl_exists(t):
        r = db_connector.execute_query(
            "SELECT COUNT(*) AS c FROM information_schema.TABLES WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME=%s",
            (t,), fetch="one"
        )
        return bool(r and int(list(r.values())[0]) > 0)

    def _cols(t):
        rows = db_connector.execute_query(
            "SELECT COLUMN_NAME FROM information_schema.COLUMNS WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME=%s",
            (t,), fetch="all"
        ) or []
        return {x["COLUMN_NAME"] for x in rows}

    def _pick(colset, *cands):
        for c in cands:
            if c and c in colset:
                return c
        return None

    def _first_existing(*tables):
        for t in tables:
            if t and _tbl_exists(t):
                return t
        return None

    # ---------- meta o produktoch ----------
    def _read_product_meta():
        meta = {}
        if not _tbl_exists("produkty"):
            return meta
        cs = _cols("produkty")
        name = _pick(cs, "nazov_vyrobku", "nazov", "produkt", "name")
        cat  = _pick(cs, "kategoria_pre_recepty", "predajna_kategoria", "kategoria", "category")
        stk  = _pick(cs, "aktualny_sklad_finalny_kg", "stav_kg", "sklad_kg", "sklad", "aktualny_sklad")
        typ  = _pick(cs, "typ_polozky", "typ_produktu", "typ", "product_type")
        pck  = _pick(cs, "vaha_balenia_g", "vaha_g", "hmotnost_g", "balenie_g")

        rows = db_connector.execute_query(
            f"SELECT {name} AS n, COALESCE({cat},'Nezaradené') AS c, "
            f"COALESCE({stk},0) AS s, COALESCE({typ},'') AS t, COALESCE({pck},0) AS g "
            f"FROM produkty", fetch="all"
        ) or []
        for r in rows:
            n = (r.get("n") or "").strip()
            if not n: continue
            t = (r.get("t") or "").upper()
            meta[n] = {
                "cat": (r.get("c") or "Nezaradené") or "Nezaradené",
                "stock": float(r.get("s") or 0),
                "is_manu": bool(t.startswith("VÝROBOK") or t.startswith("VYROBOK") or t in ("PRODUKT","PRODUCT")),
                "pack_g": float(r.get("g") or 0),
            }
        return meta

    def _stock_display(kg):
        try:
            v = float(kg or 0.0)
            return f"{int(v)} kg" if v.is_integer() else f"{v:.2f} kg"
        except Exception:
            return "—"

    product_meta = _read_product_meta()
    dates = _dates7()

    # ---------- B2B – použijeme tvoju existujúcu implementáciu, ak je dostupná ----------
    b2b = {}
    try:
        # ak máš v office_handleri pôvodnú funkciu get_7_day_order_forecast, vezmeme z nej forecast
        b2b_payload = globals().get('get_7_day_order_forecast', None)
        if callable(b2b_payload):
            res = b2b_payload()
            if isinstance(res, dict):
                b2b = res.get("forecast") or {}
    except Exception:
        b2b = {}

    # ---------- B2C – “na tvrdo” b2c_objednavky + b2c_objednavky_polozky ----------
    # Skús typické názvy (bez autodetekcie, aby to určite trafilo).
    b2c_orders_tbl = _first_existing("b2c_objednavky", "eshop_objednavky", "b2c_orders")
    b2c_items_tbl  = _first_existing("b2c_objednavky_polozky", "eshop_objednavky_polozky", "b2c_orders_items")

    b2c = {}
    b2c_debug = {"orders_tbl": b2c_orders_tbl, "items_tbl": b2c_items_tbl, "date_used": None, "rows": 0}

    if b2c_orders_tbl and b2c_items_tbl:
        oc = _cols(b2c_orders_tbl)
        ic = _cols(b2c_items_tbl)

        fk = _pick(ic, "objednavka_id", "order_id")
        name = _pick(ic, "nazov_vyrobku", "nazov", "produkt_nazov", "vyrobok_nazov", "product_name", "name")
        qty  = _pick(ic, "mnozstvo_kg", "mnozstvo", "qty_kg", "qty", "quantity")
        unit = _pick(ic, "mj", "jednotka", "unit")
        pack = _pick(ic, "vaha_balenia_g", "balenie_g", "hmotnost_g", "pack_g")
        ean  = _pick(ic, "ean_produktu", "ean")

        # dátum v orders – vyberieme prvý, ktorý existuje (poradie zodpovedá bežným schémam)
        date_candidates = ["datum_vyzdvihnutia","datum_dodania","delivery_date","pickup_date","slot_date",
                           "termin_vyzdvihnutia","termin_dodania","termin","datum","date","created_at"]
        date_cols = [c for c in date_candidates if c in oc]
        if date_cols:
            date_expr = "DATE(COALESCE(" + ", ".join("o."+c for c in date_cols) + "))"
        else:
            # ak nemáš nič v objednávke, posledný pokus – položka (stáva sa pri niektorých eshop exportoch)
            date_i_candidates = ["datum_vyzdvihnutia","datum_dodania","delivery_date","pickup_date","slot_date",
                                 "termin_vyzdvihnutia","termin_dodania","termin","datum","date","created_at"]
            date_i_cols = [c for c in date_i_candidates if c in ic]
            date_expr = "DATE(COALESCE(" + ", ".join("pol."+c for c in date_i_cols) + "))" if date_i_cols else "DATE(o.created_at)"
        b2c_debug["date_used"] = date_expr

        status = _pick(oc, "stav", "status")
        where = [f"{date_expr} BETWEEN %s AND %s"]
        params = (dates[0], dates[-1])
        if status:
            where.append(
                f"COALESCE(CONVERT(o.{status} USING utf8mb4) COLLATE {COLL}, '') NOT IN "
                "('Zrušená','Zrusena','Zrušena','Zrušené','Cancelled')"
            )

        pack_expr = f"pol.{pack}" if pack else "p.vaha_balenia_g"

        join = f"""
            JOIN {b2c_items_tbl} pol ON pol.{fk} = o.id
            LEFT JOIN produkty p ON (
                {"(p.ean IS NOT NULL AND pol."+ean+" IS NOT NULL AND CONVERT(p.ean USING utf8mb4) COLLATE "+COLL+" = CONVERT(pol."+ean+" USING utf8mb4) COLLATE "+COLL+") OR" if ean else ""}
                (CONVERT(p.nazov_vyrobku USING utf8mb4) COLLATE {COLL} = CONVERT(pol.{name} USING utf8mb4) COLLATE {COLL})
            )
        """

        sql = f"""
            SELECT
              COALESCE(CONVERT(pol.{name} USING utf8mb4) COLLATE {COLL}, '') AS n,
              {date_expr} AS d,
              SUM(
                CASE
                  WHEN LOWER(COALESCE(pol.{unit},'')) = 'kg' THEN COALESCE(pol.{qty},0)
                  WHEN LOWER(COALESCE(pol.{unit},'')) = 'g'  THEN COALESCE(pol.{qty},0) / 1000
                  WHEN LOWER(COALESCE(pol.{unit},'')) IN ('ks','pc','pcs') THEN COALESCE(pol.{qty},0) * COALESCE({pack_expr}, 0) / 1000
                  ELSE COALESCE(pol.{qty},0)
                END
              ) AS q
            FROM {b2c_orders_tbl} o
            {join}
            WHERE {" AND ".join(where)}
            GROUP BY n, d
            ORDER BY n, d
        """
        rows = db_connector.execute_query(sql, params, fetch="all") or []
        b2c_debug["rows"] = len(rows)

        idx = {}
        for r in rows:
            n = (r.get("n") or "").strip()
            d = (r.get("d") or "")
            if not n or d not in dates: continue
            q = float(r.get("q") or 0.0)

            meta = product_meta.get(n, {"cat":"Nezaradené","stock":0.0,"is_manu":True})
            cat = meta["cat"] or "Nezaradené"
            key = (cat, n)

            if key not in idx:
                item = {
                    "name": n,
                    "mj": "kg",
                    "stock_display": _stock_display(meta["stock"]),
                    "isManufacturable": bool(meta["is_manu"]),
                    "daily_needs": {dt: 0 for dt in dates},
                }
                b2c.setdefault(cat, []).append(item)
                idx[key] = item
            idx[key]["daily_needs"][d] = idx[key]["daily_needs"].get(d, 0) + q

        for cat, arr in b2c.items():
            def _total(it): return sum(float(it["daily_needs"].get(dt,0) or 0) for dt in dates)
            arr.sort(key=lambda it: (-_total(it), it["name"]))

    # ---------- MERGE B2B + B2C do forecast ----------
    def _merge_two(fa, fb):
        out = {}
        for src in (fa or {}), (fb or {}):
            for cat, items in (src or {}).items():
                out.setdefault(cat, [])
                idx = {(p["name"], p.get("mj","kg")): i for i,p in enumerate(out[cat])}
                for p in items or []:
                    k = (p["name"], p.get("mj","kg"))
                    if k in idx:
                        tgt = out[cat][idx[k]]
                    else:
                        tgt = {
                            "name": p["name"],
                            "mj": p.get("mj","kg"),
                            "stock_display": p.get("stock_display","—"),
                            "isManufacturable": bool(p.get("isManufacturable",True)),
                            "daily_needs": {dt: 0 for dt in p.get("daily_needs",{}).keys() or []}
                        }
                        if not tgt["daily_needs"]:
                            for dt in (p.get("daily_needs") or {}).keys():
                                tgt["daily_needs"][dt] = 0
                        out[cat].append(tgt); idx[k] = len(out[cat]) - 1
                    for dt, val in (p.get("daily_needs") or {}).items():
                        tgt["daily_needs"][dt] = float(tgt["daily_needs"].get(dt,0)) + float(val or 0)
                    if len(p.get("stock_display","")) > len(tgt.get("stock_display","")):
                        tgt["stock_display"] = p.get("stock_display","—")
                    tgt["isManufacturable"] = bool(tgt.get("isManufacturable",True) or p.get("isManufacturable",True))
        for cat, arr in out.items():
            def _total(it): return sum(float(v or 0) for v in it["daily_needs"].values())
            arr.sort(key=lambda it: (-_total(it), it["name"]))
        return out

    merged = _merge_two(b2b, b2c)

    return {
        "dates": dates,
        "forecast": merged,           # HLAVNÉ pole – B2B+B2C DOKOPY
        "b2c_forecast": b2c or {},    # čisté B2C
        "forecast_b2c": b2c or {},    # alias
        "b2c": b2c or {},             # alias
        "debug": {
            "b2c": b2c_debug,
            "b2b_present": bool(b2b),
        }
    }
